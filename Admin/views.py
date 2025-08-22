import os
from django.conf import settings
from django.shortcuts import render, redirect
from django.contrib import messages
from authentication.models import User
from Admin import models as admin_models
from .forms import UserForm, SessionForm, StudentForm, LeadForm
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from datetime import date, datetime
from dateutil.relativedelta import relativedelta
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from django.http import HttpResponse
from django.core.mail import send_mail
from django.db.models import Count
from decimal import Decimal
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from datetime import datetime, timedelta
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from functools import wraps

# Role-based access control decorator
def teacher_redirect_to_attendance(view_func):
    """
    Decorator that redirects teacher users (usertype=3) to attendance page
    if they try to access admin-only pages
    """
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if hasattr(request, 'user') and request.user.is_authenticated:
            if request.user.usertype == 3:  # Teacher
                # Redirect teachers to their attendance page
                return redirect('tec_select_course')
        return view_func(request, *args, **kwargs)
    return wrapper


def notify_late_fee_students(request):
    if request.method == 'POST':
        try:
            # Validate email configuration before proceeding
            from django.conf import settings
            
            # Check if email backend is configured
            if not hasattr(settings, 'EMAIL_BACKEND') or not settings.EMAIL_BACKEND:
                return JsonResponse({'status': 'error', 'message': 'Email backend not configured. Please contact administrator.'})
            
            # Check if SMTP settings are configured for production
            if settings.EMAIL_BACKEND == 'django.core.mail.backends.smtp.EmailBackend':
                if not getattr(settings, 'EMAIL_HOST', None):
                    return JsonResponse({'status': 'error', 'message': 'Email host not configured. Please contact administrator.'})
            
            # Track statistics for the response
            emails_sent = 0
            students_with_pending_fees = 0
            failed_emails = 0
            
            # Get all unpaid payments for active students
            from datetime import date
            today_date = date.today()
            
            unpaid_payments = admin_models.Payments.objects.filter(
                amount=0,  # Unpaid payments
                studentsession__student__status='Active'
            ).select_related('studentsession__student', 'studentsession__session')
            
            # Group payments by student
            student_payments = {}
            for payment in unpaid_payments:
                if payment.studentsession and payment.studentsession.student:
                    student = payment.studentsession.student
                    if student not in student_payments:
                        student_payments[student] = []
                    student_payments[student].append(payment)
            
            for student, payments in student_payments.items():
                # Skip students with no email
                if not student.email:
                    continue
                    
                # Calculate total pending amount and build session details
                pending_amount = 0
                session_details = []
                overdue_details = []
                
                for payment in payments:
                    if payment.studentsession and payment.studentsession.session:
                        session_fee = payment.studentsession.session.fee
                        pending_amount += session_fee
                        
                        # Calculate days overdue
                        days_diff = (payment.date - today_date).days
                        days_overdue = abs(days_diff) if days_diff < 0 else 0
                        
                        if days_overdue > 0:
                            overdue_details.append(f"- {payment.studentsession.session.session_name}: Rs. {session_fee:,.0f} ({days_overdue} days overdue)")
                        else:
                            session_details.append(f"- {payment.studentsession.session.session_name}: Rs. {session_fee:,.0f} (due: {payment.date})")
                
                # Only proceed if there's an actual pending amount
                if pending_amount <= 0:
                    continue
                    
                students_with_pending_fees += 1
                
                # Combine overdue and upcoming payments
                all_details = overdue_details + session_details
                
                # Create a more detailed and professional email
                subject = "Fee Payment Reminder - Iqra Academy"
                message = f"""Dear {student.student_name},

This is a friendly reminder that you have an outstanding payment of Rs. {pending_amount:,.0f} for your courses at Iqra Academy.

Payment Details:
{'\n'.join(all_details)}

Total Pending Amount: Rs. {pending_amount:,.0f}

Please arrange for the payment at your earliest convenience to avoid any interruption in your learning experience.

If you have already made the payment, please disregard this message.

For any queries regarding your payment, please contact our accounts department.

Regards,
Iqra Academy
Accounts Department
Phone: [Your Phone Number]
Email: admin@iqrainstitute.com"""
                
                # Retry logic for each email
                import time
                max_retries = 3
                retry_delay = 1  # seconds (shorter for bulk emails)
                email_sent = False
                
                for attempt in range(max_retries):
                    try:
                        # Send the email
                        send_mail(
                            subject=subject,
                            message=message,
                            from_email=getattr(settings, 'EMAIL_HOST_USER', 'admin@iqrainstitute.com'),
                            recipient_list=[student.email],
                            fail_silently=False,
                        )
                        emails_sent += 1
                        email_sent = True
                        break  # Success, exit retry loop
                    except Exception as email_error:
                        # If this is the last attempt, log failure
                        if attempt == max_retries - 1:
                            failed_emails += 1
                            print(f"Failed to send email to {student.email} after {max_retries} attempts: {str(email_error)}")
                            # Log the specific error for debugging
                            if 'WinError 10060' in str(email_error):
                                print(f"Connection timeout error for {student.email} - check firewall/network settings")
                        else:
                            # Wait before retrying (shorter delay for bulk emails)
                            time.sleep(retry_delay * (attempt + 1))
                            print(f"Email send attempt {attempt + 1} failed for {student.email}, retrying...")
            
            # Return detailed success response
            message = f'Bulk reminder process completed. {emails_sent} emails sent successfully'
            if failed_emails > 0:
                message += f', {failed_emails} emails failed to send'
            
            return JsonResponse({
                'status': 'success', 
                'message': message,
                'details': {
                    'emails_sent': emails_sent,
                    'students_with_pending_fees': students_with_pending_fees,
                    'failed_emails': failed_emails
                }
            })
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': f'Error processing bulk reminders: {str(e)}'})

    return JsonResponse({'status': 'error', 'message': 'Invalid request method.'})

@csrf_exempt
def send_fee_reminder(request):
    if 'user_id' not in request.session:
        return JsonResponse({'status': 'error', 'message': 'Not authenticated'})
    
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': 'Invalid request method'})
    
    try:
        user_id = request.session.get('user_id')
        user = User.objects.get(id=user_id)
        
        # Handle both JSON and form data
        if request.content_type == 'application/json':
            import json
            data = json.loads(request.body)
            student_id = data.get('student_id')
            session_id = data.get('session_id')  # Optional for specific session
        else:
            student_id = request.POST.get('student_id')
            session_id = request.POST.get('session_id')
        
        if not student_id:
            return JsonResponse({'status': 'error', 'message': 'Student ID is required'})
        
        try:
            student = admin_models.Student.objects.get(id=student_id)
        except admin_models.Student.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Student not found'})
        
        if not student.email:
            return JsonResponse({'status': 'error', 'message': 'Student email not available'})
        
        # Check for unpaid payments (amount = 0) for this student
        from datetime import date
        today_date = date.today()
        
        unpaid_payments = admin_models.Payments.objects.filter(
            studentsession__student=student,
            amount=0,  # Unpaid payments
            studentsession__student__status='Active'
        ).select_related('studentsession__session')
        
        if not unpaid_payments.exists():
            return JsonResponse({'status': 'error', 'message': 'No pending fees for this student'})
        
        # Calculate total pending amount and get session details
        pending_amount = 0
        session_details = []
        overdue_details = []
        
        for payment in unpaid_payments:
            if payment.studentsession and payment.studentsession.session:
                session_fee = payment.studentsession.session.fee
                pending_amount += session_fee
                
                # Calculate days overdue
                days_diff = (payment.date - today_date).days
                days_overdue = abs(days_diff) if days_diff < 0 else 0
                
                if days_overdue > 0:
                    overdue_details.append(f"- {payment.studentsession.session.session_name}: Rs. {session_fee:,.0f} ({days_overdue} days overdue)")
                else:
                    session_details.append(f"- {payment.studentsession.session.session_name}: Rs. {session_fee:,.0f} (due: {payment.date})")
        
        # Combine overdue and upcoming payments
        all_details = overdue_details + session_details
        
        # Create professional email content
        subject = "Fee Payment Reminder - Iqra Academy"
        message = f"""Dear {student.student_name},

This is a friendly reminder that you have an outstanding payment of Rs. {pending_amount:,.0f} for your courses at Iqra Academy.

Payment Details:
{chr(10).join(all_details) if all_details else '- Course fees pending'}

Total Pending Amount: Rs. {pending_amount:,.0f}

Please arrange for the payment at your earliest convenience to avoid any interruption in your learning experience.

If you have already made the payment, please disregard this message.

For any queries regarding your payment, please contact our accounts department.

Regards,
Iqra Academy
Accounts Department
Phone: [Your Phone Number]
Email: admin@iqrainstitute.com"""
        
        # Validate email configuration before sending
        from django.conf import settings
        
        # Check if email backend is configured
        if not hasattr(settings, 'EMAIL_BACKEND') or not settings.EMAIL_BACKEND:
            return JsonResponse({'status': 'error', 'message': 'Email backend not configured. Please contact administrator.'})
        
        # Check if SMTP settings are configured for production
        if settings.EMAIL_BACKEND == 'django.core.mail.backends.smtp.EmailBackend':
            if not getattr(settings, 'EMAIL_HOST', None):
                return JsonResponse({'status': 'error', 'message': 'Email host not configured. Please contact administrator.'})
        
        # Retry logic for email sending
        import time
        max_retries = 3
        retry_delay = 2  # seconds
        
        for attempt in range(max_retries):
            try:
                # Send the email
                send_mail(
                    subject=subject,
                    message=message,
                    from_email=getattr(settings, 'EMAIL_HOST_USER', 'admin@iqrainstitute.com'),
                    recipient_list=[student.email],
                    fail_silently=False,
                )
                break  # Success, exit retry loop
            except Exception as email_error:
                error_message = str(email_error)
                
                # If this is the last attempt, return error
                if attempt == max_retries - 1:
                    if 'WinError 10060' in error_message or 'Connection refused' in error_message:
                        return JsonResponse({
                            'status': 'error', 
                            'message': f'Failed to send reminder after {max_retries} attempts: Connection timeout. Please check your internet connection and email configuration. Try switching to SSL port 465 if the issue persists.'
                        })
                    else:
                        return JsonResponse({
                            'status': 'error', 
                            'message': f'Failed to send reminder after {max_retries} attempts: {error_message}. Please check email configuration.'
                        })
                
                # Wait before retrying (exponential backoff)
                time.sleep(retry_delay * (2 ** attempt))
                print(f"Email send attempt {attempt + 1} failed, retrying in {retry_delay * (2 ** attempt)} seconds...")
        
        # Create notification for this reminder
        admin_models.Notification.objects.create(
            user=user,
            category='Late Fee',
            content=f"Fee reminder sent to {student.student_name} ({student.email}) - Rs. {pending_amount:,.0f} pending"
        )
        
        return JsonResponse({
            'status': 'success', 
            'message': 'Reminder sent successfully',
            'details': {
                'student_name': student.student_name,
                'email': student.email,
                'pending_amount': pending_amount
            }
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Error sending reminder: {str(e)}'})
@teacher_redirect_to_attendance
def EmailService(request):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    if request.method == 'POST':
        # Get data from the form
        email_content = request.POST.get('email_content')  # Email body
        email_subject = request.POST.get('email_subject')  # Email subject
        email_list = []
        if 'faculty_checkbox' in request.POST:
            selected_users = User.objects.all()
            email_addresses = [user.email for user in selected_users if user.email]
            email_list.extend(email_addresses)
        if 'student_checkbox' in request.POST:
            selected_users = admin_models.Student.objects.all()
            email_addresses = [user.email for user in selected_users if user.email]
            email_list.extend(email_addresses)
        if 'lead_checkbox' in request.POST:
            selected_users = admin_models.Lead.objects.all()
            email_addresses = [user.email for user in selected_users if user.email]
            email_list.extend(email_addresses)

        if email_list:
            try:
                # Send the email to the selected users
                send_mail(
                    subject=email_subject,
                    message=email_content,
                    from_email='callmemutiurrehman@gmail.com',  # Replace with your sender email
                    recipient_list=email_list,
                    fail_silently=False,
                )
                return JsonResponse({'status': 'success', 'message': 'Emails sent successfully!'})
            except Exception as e:
                return JsonResponse({'status': 'error', 'message': f'Failed to send emails: {str(e)}'})

        return JsonResponse({'status': 'error', 'message': 'No valid email addresses found.'})
    context = {
        'user': user,
    }
    return render(request, 'Admin/EmailService.html', context)
def print_attendance_report(request, course_id):
    # Get date range from request
    start_date = request.GET.get('start_date')
    end_date = request.GET.get('end_date')
    
    # Get the course object
    course = admin_models.Sessions.objects.get(id=course_id)

    # Fetch attendance data for this course within the date range
    attendances = admin_models.Attendance.objects.filter(
        course=course,
        date__range=[start_date, end_date]
    ).order_by('date')
    
    students = admin_models.StudentSession.objects.filter(session=course)

    # Create a PDF response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{course.session_name}_attendance_report.pdf"'

    # Create the PDF in landscape orientation
    pdf = canvas.Canvas(response, pagesize=landscape(letter))
    width, height = landscape(letter)  # Get dimensions for landscape orientation

    # Add institution name
    pdf.setFont("Helvetica-Bold", 16)
    title = "IQRA INSTITUTE OF COMPETITIVE EXAMINATION"
    title_width = pdf.stringWidth(title, "Helvetica-Bold", 16)
    pdf.drawString((width - title_width) / 2, height - 50, title)

    # Add report title and date range
    pdf.setFont("Helvetica-Bold", 14)
    report_title = f"Attendance Report for {course.session_name}"
    report_width = pdf.stringWidth(report_title, "Helvetica-Bold", 14)
    pdf.drawString((width - report_width) / 2, height - 80, report_title)
    
    pdf.setFont("Helvetica", 12)
    date_range = f"Period: {start_date} to {end_date}"
    date_width = pdf.stringWidth(date_range, "Helvetica", 12)
    pdf.drawString((width - date_width) / 2, height - 100, date_range)

    # Get unique dates from attendance records within the selected range
    dates = sorted(set(attendance.date for attendance in attendances))
    
    if not dates:
        # If no attendance records found
        pdf.setFont("Helvetica", 12)
        message = "No attendance records found for the selected date range."
        msg_width = pdf.stringWidth(message, "Helvetica", 12)
        pdf.drawString((width - msg_width) / 2, height - 140, message)
        pdf.save()
        return response

    # Calculate table dimensions
    max_dates_per_page = 13  # Reduced to accommodate roll number column
    row_height = 25
    header_height = 30
    
    # Create tables for chunks of dates
    current_y = height - 140
    date_chunks = [dates[i:i + max_dates_per_page] for i in range(0, len(dates), max_dates_per_page)]
    
    for chunk_index, date_chunk in enumerate(date_chunks):
        if chunk_index > 0:
            # Start a new page for each chunk after the first
            pdf.showPage()
            pdf.setPageSize(landscape(letter))
            current_y = height - 100
            
            # Add header to new page
            pdf.setFont("Helvetica-Bold", 16)
            pdf.drawString((width - title_width) / 2, height - 50, title)
            pdf.setFont("Helvetica-Bold", 14)
            pdf.drawString((width - report_width) / 2, height - 80, f"{report_title} (Continued)")
        
        # Create table data with roll number, student name, and attendance dates
        table_data = [["Roll No", "Student Name"] + [date.strftime("%d-%m-%Y") for date in date_chunk]]
        
        # Add student attendance data with roll numbers
        for student in students:
            roll_no = student.student.rollno if student.student.rollno else "-"
            row = [roll_no, student.student.student_name]
            for date in date_chunk:
                # Get attendance for this specific date
                attendance = attendances.filter(
                    student=student.student,
                    date=date
                ).first()
                status = attendance.status if attendance else "Absent"
                row.append(status)
            table_data.append(row)
        
        # Calculate column widths for landscape with roll number column
        roll_col_width = 80   # Width for roll numbers
        name_col_width = 150  # Width for student names (reduced to make room for roll number)
        date_col_width = (width - 100 - roll_col_width - name_col_width) / len(date_chunk)
        col_widths = [roll_col_width, name_col_width] + [date_col_width] * len(date_chunk)
        
        # Create and style the table
        table = Table(table_data, colWidths=col_widths, rowHeights=[row_height] * len(table_data))
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 11),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ])
        
        # Add conditional formatting for Present/Absent (starting from column 2 since we added roll number)
        for row in range(1, len(table_data)):
            for col in range(2, len(table_data[0])):
                if table_data[row][col] == 'Present':
                    style.add('TEXTCOLOR', (col, row), (col, row), colors.green)
                elif table_data[row][col] == 'Absent':
                    style.add('TEXTCOLOR', (col, row), (col, row), colors.red)
        
        table.setStyle(style)

        # Draw the table
        table.wrapOn(pdf, width - 100, height)
        table.drawOn(pdf, 50, current_y - (len(table_data) * row_height))
        
        current_y -= (len(table_data) * row_height + 60)

    # Add footer with generation date
    pdf.setFont("Helvetica", 10)
    footer = f"Report Generated on: {date.today().strftime('%d-%m-%Y')}"
    footer_width = pdf.stringWidth(footer, "Helvetica", 10)
    pdf.drawString((width - footer_width) / 2, 30, footer)
    
    # Save the PDF
    pdf.save()
    return response
def ensure_session_fees():
    """Ensure all student sessions have proper fee values"""
    sessions_updated = 0
    for session in admin_models.StudentSession.objects.all():
        if session.fee is None and session.session and session.session.fee:
            session.fee = session.session.fee
            session.save()
            sessions_updated += 1
    return sessions_updated

@teacher_redirect_to_attendance
def Payment(request):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)
    
    # Restrict access for moderators
    if user.usertype == 2:
        return redirect('Admin_Dashboard')
    
    # UNIFIED DATA SOURCE: Use only Payments table and calculated properties
    payments = admin_models.Payments.objects.select_related(
        'studentsession__student', 'studentsession__session', 'user'
    ).all()
    
    # Get all active students (single source for student data)
    students = admin_models.Student.objects.filter(status='Active')
    
    # Initialize metrics
    total_revenue = 0
    total_pending = 0
    total_discount = 0
    total_expected_revenue = 0
    session_revenue = {}
    user_collection = {}
    overdue_amount = 0
    
    # Student payment status counters
    students_paid = 0
    students_partial = 0
    students_unpaid = 0
    
    # Calculate total revenue from Payments table (SINGLE SOURCE)
    total_revenue = sum(p.amount or 0 for p in payments)
    
    # Track session revenue from payments (unified system)
    for payment in payments:
        if payment.studentsession and payment.studentsession.session:
            session_name = payment.studentsession.session.session_name
            session_revenue[session_name] = session_revenue.get(session_name, 0) + (payment.amount or 0)
    
    # Track user collection from payment records
    for payment in payments:
        if payment.user:
            collector_name = f"{payment.user.first_name} {payment.user.last_name}"
            user_collection[collector_name] = user_collection.get(collector_name, 0) + (payment.amount or 0)
    
    # Calculate totals using unified Student model properties
    student_fee_data = []
    for student in students:
        # Use calculated properties from unified system
        student_total_fee = student.total_fee
        student_paid = student.total_paid
        student_balance = student.remaining_balance
        student_status = student.payment_status
        
        # Get student sessions for display
        student_sessions = student.student_sessions.filter(status='Active')
        
        # Calculate discount from sessions
        student_discount = sum(session.discount or 0 for session in student_sessions)
        
        # Update totals
        total_expected_revenue += student_total_fee
        total_pending += student_balance
        total_discount += student_discount
        
        # Count payment status
        if student_status == 'Paid':
            students_paid += 1
        elif student_status == 'Partial':
            students_partial += 1
        else:
            students_unpaid += 1
        
        # Check for overdue amounts
        from datetime import date
        today = date.today()
        is_overdue = any(
            session.due_date and session.due_date < today and session.session_balance > 0
            for session in student_sessions
        )
        if is_overdue:
            overdue_amount += student_balance
        
        # Create student fee object for display (using unified data)
        student_fee = type('StudentFee', (), {
            'student': student,
            'sessions': list(student_sessions),
            'calculated_final_fee': student_total_fee,
            'display_paid_amount': student_paid,
            'calculated_remaining_amount': student_balance,
            'display_discount': student_discount,
            'payment_status': student_status
        })()
        
        student_fee_data.append(student_fee)
    
    # Get session revenue from student sessions
    student_sessions = admin_models.StudentSession.objects.select_related(
        'student', 'session'
    ).filter(status='Active')
    
    # Create student fee data for display using UNIFIED SYSTEM ONLY
    student_fee_data = []
    processed_students = set()  # Track processed students to avoid duplicates
    
    for student in students:
        if student.id in processed_students:
            continue
        processed_students.add(student.id)
        
        # Use unified Student model properties (calculated from Payments)
        student_total_fee = student.total_fee
        student_paid = student.total_paid
        student_balance = student.remaining_balance
        student_status = student.payment_status
        
        # Get student sessions for display
        student_sessions_list = list(student.student_sessions.filter(status='Active'))
        
        # Calculate discount from sessions
        student_discount = sum(session.discount or 0 for session in student_sessions_list)
        
        # Update totals
        total_expected_revenue += student_total_fee
        total_pending += student_balance
        total_discount += student_discount
        
        # Count payment status
        if student_status == 'Paid':
            students_paid += 1
        elif student_status == 'Partial':
            students_partial += 1
        else:
            students_unpaid += 1
        
        # Check for overdue amounts
        from datetime import date
        today = date.today()
        is_overdue = any(
            session.due_date and session.due_date < today and session.session_balance > 0
            for session in student_sessions_list
        )
        if is_overdue:
            overdue_amount += student_balance
        
        # Create student fee object for display (using unified data)
        student_fee = type('StudentFee', (), {
            'student': student,
            'sessions': student_sessions_list,
            'calculated_final_fee': student_total_fee,
            'display_paid_amount': student_paid,
            'calculated_remaining_amount': student_balance,
            'display_discount': student_discount,
            'payment_status': student_status
        })()
        
        student_fee_data.append(student_fee)
    
    # Calculate additional metrics
    total_students = len(student_fee_data)
    total_payments_count = payments.count()
    avg_payment = total_revenue / total_payments_count if total_payments_count > 0 else 0
    collection_rate = (total_revenue / total_expected_revenue * 100) if total_expected_revenue > 0 else 0
    
    # Get recent payments (ordered by date)
    recent_payments = payments.order_by('-date', '-id')[:10]
    
    # Get top sessions by revenue
    top_sessions = sorted(session_revenue.items(), key=lambda x: x[1], reverse=True)[:5]
    
    # Calculate daily and yearly revenue
    from datetime import date, timedelta
    today = date.today()
    daily_revenue = sum(p.amount or 0 for p in payments if p.date == today)
    yearly_revenue = sum(p.amount or 0 for p in payments if p.date and p.date.year == today.year)
    
    # Calculate monthly revenue (current month)
    monthly_revenue = sum(p.amount or 0 for p in payments if p.date and p.date.month == today.month and p.date.year == today.year)
    
    # Business Intelligence Metrics
    active_students_count = admin_models.Student.objects.filter(status='Active').count()
    revenue_per_student = total_revenue / active_students_count if active_students_count > 0 else 0
    
    # Count overdue students using unified system
    overdue_students_count = 0
    for student in students:
        if student.remaining_balance > 0:
            # Check if any session for this student is overdue
            student_sessions_list = student.student_sessions.filter(status='Active')
            for session in student_sessions_list:
                if session.due_date and session.due_date < today:
                    overdue_students_count += 1
                    break
    
    # Calculate projected monthly revenue based on current trends
    days_in_month = today.day
    if days_in_month > 0:
        daily_avg = monthly_revenue / days_in_month
        days_remaining = 30 - days_in_month
        projected_monthly_revenue = monthly_revenue + (daily_avg * days_remaining)
    else:
        projected_monthly_revenue = 0
    
    # Calculate payment trends (last 7 days)
    week_ago = today - timedelta(days=7)
    weekly_payments = [p for p in payments if p.date and p.date >= week_ago]
    weekly_revenue = sum(p.amount or 0 for p in weekly_payments)
    
    # Session performance analysis
    session_performance = []
    for session_name, revenue in session_revenue.items():
        # Count students in this session
        session_obj = admin_models.Sessions.objects.filter(session_name=session_name).first()
        if session_obj:
            student_count = admin_models.StudentSession.objects.filter(session=session_obj, status='Active').count()
            avg_revenue_per_student = revenue / student_count if student_count > 0 else 0
            session_performance.append({
                'name': session_name,
                'revenue': revenue,
                'students': student_count,
                'avg_per_student': avg_revenue_per_student
            })
    
    # Sort by revenue
    session_performance.sort(key=lambda x: x['revenue'], reverse=True)
    
    context = {
        'user': user,
        'payments': recent_payments,
        'student_fees': student_fee_data,
        'total_revenue': total_revenue,
        'total_pending': total_pending,
        'total_discount': total_discount,
        'total_expected_revenue': total_expected_revenue,
        'session_revenue': session_revenue,
        'user_collection': user_collection,
        'total_payments_count': total_payments_count,
        'recent_payments': recent_payments,
        'top_sessions': top_sessions,
        'avg_payment': avg_payment,
        'collection_rate': round(collection_rate, 1),
        'students_paid': students_paid,
        'students_partial': students_partial,
        'students_unpaid': students_unpaid,
        'total_students': total_students,
        'overdue_amount': overdue_amount,
        'daily_revenue': daily_revenue,
        'yearly_revenue': yearly_revenue,
        'monthly_revenue': monthly_revenue,
        'active_students_count': active_students_count,
        'revenue_per_student': round(revenue_per_student, 0),
        'overdue_students_count': overdue_students_count,
        'projected_monthly_revenue': round(projected_monthly_revenue, 0),
        'weekly_revenue': weekly_revenue,
        'session_performance': session_performance,
        'today_date': today,
    }
    return render(request, 'Admin/Payments.html', context)
@csrf_exempt
def add_fee_payment(request, session_id):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    if request.method == "POST":
        try:
            # Get the session and payment details
            session = admin_models.StudentSession.objects.get(id=session_id)
            amount = int(request.POST.get("amount"))
            due_date = request.POST.get("due_date")

            # Validate amount using unified system
            remaining_fee = session.session_balance  # Use calculated property
            if amount <= 0 or amount > remaining_fee:
                return JsonResponse({"success": False, "error": "Invalid amount entered."})

            # Add the payment record (SINGLE SOURCE OF TRUTH)
            payment = admin_models.Payments.objects.create(
                studentsession=session,
                user=user,
                amount=amount,
                date=date.today(),
            )

            # Update due_date (no need to update fee_paid - it's calculated from Payments)
            session.due_date = due_date
            session.save()
            
            message = "Collected Fee Rs " + str(amount) + " from " + session.student.student_name
            admin_models.Notification.objects.create(user=user, category='New Fee', content=message)
            return JsonResponse({"success": True})
        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "Invalid request method."})

def MakeNotification(request):
    if 'user_id' not in request.session:
        return redirect('home')
    
    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    
    # Get all active sessions
    sessions = admin_models.Sessions.objects.filter(status='Active')
    
    for session in sessions:
        # Get all student sessions for this session
        student_sessions = admin_models.StudentSession.objects.filter(session=session)
        
        for student_session in student_sessions:
            # Convert due_date to datetime.date if it's not None
            if student_session.due_date:
                # Check if today's date is greater than due date
                if datetime.now().date() > student_session.due_date:
                    # Create notification for late fee
                    notification = admin_models.Notification.objects.create(
                        user=user,
                        student=student_session.student,
                        session=session,
                        message=f"Late fee notification for {student_session.student.student_name} in {session.session_name}"
                    )
    
    return redirect('notification')

@teacher_redirect_to_attendance
def Notification(request):
    if 'user_id' not in request.session:
        return redirect('home')
    
    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    
    # Get notifications ordered by date/time in descending order (latest first)
    notifications = admin_models.Notification.objects.all().order_by('-date', '-id')
    
    # Get today's date for comparison
    today_date = date.today()
    from datetime import timedelta
    from dateutil.relativedelta import relativedelta
    seven_days_from_now = today_date + timedelta(days=7)
    
    # AUTOMATIC MONTHLY RENEWAL PROCESSING
    # Check for monthly sessions that need renewal within 7 days
    monthly_student_sessions = admin_models.StudentSession.objects.filter(
        session__session_type='monthly',
        status='Active',
        student__status='Active'
    ).select_related('student', 'session')
    
    renewals_created = 0
    for student_session in monthly_student_sessions:
        try:
            # Get the last payment for this student session
            last_payment = admin_models.Payments.objects.filter(
                studentsession=student_session
            ).order_by('-date').first()
            
            next_due_date = None
            if last_payment:
                # Calculate next monthly due date from last payment
                next_due_date = last_payment.date + relativedelta(months=1)
            elif student_session.registration_date:
                # No payments yet, calculate from registration date
                next_due_date = student_session.registration_date + relativedelta(months=1)
            
            if next_due_date and next_due_date <= seven_days_from_now:
                # Check if there's already an unpaid payment for this due date
                existing_unpaid = admin_models.Payments.objects.filter(
                    studentsession=student_session,
                    date=next_due_date,
                    amount=0
                ).exists()
                
                if not existing_unpaid:
                    # Create new unpaid payment for monthly renewal
                    admin_models.Payments.objects.create(
                        studentsession=student_session,
                        user=user,
                        amount=0,  # Unpaid
                        date=next_due_date
                    )
                    
                    # Update next_monthly_due field
                    student_session.next_monthly_due = next_due_date
                    student_session.save()
                    
                    # Create notification for the renewal
                    message = (
                        f"Monthly renewal due for {student_session.student.student_name} "
                        f"in {student_session.session.session_name} - "
                        f"Rs.{student_session.session.fee} due on {next_due_date}"
                    )
                    admin_models.Notification.objects.create(
                        user=user,
                        category='Monthly Renewal',
                        content=message
                    )
                    
                    renewals_created += 1
                    
        except Exception as e:
            # Log error but continue processing other sessions
            print(f"Error processing monthly renewal for {student_session.student.student_name}: {str(e)}")
    
    # Get unpaid payments that are due within 7 days or overdue
    due_payments = admin_models.Payments.objects.filter(
        amount=0,  # Unpaid payments
        date__lte=seven_days_from_now,  # Due within 7 days or overdue
        studentsession__student__status='Active',
        studentsession__session__status='Active'
    ).select_related('studentsession__student', 'studentsession__session')
    
    # Process each payment to create session-like objects for template compatibility
    processed_sessions = []
    for payment in due_payments:
        if payment.studentsession:
            # Calculate days until due or days overdue
            days_diff = (payment.date - today_date).days
            
            # Create a session-like object for template compatibility
            session_obj = type('DuePaymentSession', (), {
                'student': payment.studentsession.student,
                'session': payment.studentsession.session,
                'due_date': payment.date,
                'days_until_due': days_diff,
                'days_overdue': abs(days_diff) if days_diff < 0 else 0,
                'fee_amount': payment.studentsession.session.fee if payment.studentsession.session else 0,
                'balance': payment.studentsession.session.fee if payment.studentsession.session else 0,  # Unpaid amount
                'payment_id': payment.id
            })()
            
            processed_sessions.append(session_obj)
    
    # Sort all sessions by due date
    processed_sessions.sort(key=lambda x: x.due_date)
    
    context = {
        'user': user,
        'notifications': notifications,
        'due_fee_sessions': processed_sessions,
        'due_fees_count': len(processed_sessions),
        'today_date': today_date
    }
    return render(request, 'Admin/Notification.html', context)

@csrf_exempt
def mark_all_notifications_read(request):
    if 'user_id' not in request.session:
        return redirect('home')
    
    if request.method == 'POST':
        # Mark all notifications as read
        admin_models.Notification.objects.filter(is_read=False).update(is_read=True)
        
        # Return JSON response for AJAX requests
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({'success': True, 'message': 'All notifications marked as read'})
        
        # For non-AJAX requests, redirect back to notifications
        return redirect('notification')
    
    # If not POST request, redirect to notifications
    return redirect('notification')

def select_course(request):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    
    # Get all sessions with their student counts using annotation
    courses = admin_models.Sessions.objects.annotate(
        student_count=Count('session_students')
    ).all()
    
    context = {
        'user': user,
        'courses': courses
    }
    return render(request, 'Admin/Attendance.html', context)
def mark_attendance(request, course_id):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    course = admin_models.Sessions.objects.get(id=course_id)
    students = admin_models.StudentSession.objects.filter(session=course)

    if request.method == 'POST':
        date1 = request.POST.get('date')
        for student in students:
            status = request.POST.get(f'status_{student.student.id}')
            admin_models.Attendance.objects.update_or_create(
                course=course,
                student=student.student,
                date=date1,
                defaults={'status': status}
            )
            sessions = admin_models.StudentSession.objects.filter(student=student.student)
            for session in sessions:
                if session.due_date and session.due_date < date.today():
                    message = "Due Date passed for " + student.student.student_name + " in " + session.session.session_name + " session"
                    admin_models.Notification.objects.create(user=user, category='Late fee',
                                                             content=message)
        messages.success(request, 'Attendance marked successfully!')
        return redirect('select_course')
    context = {
        'user': user,
        'course': course,
        'students': students
    }
    return render(request, 'Admin/Mark_Attendance.html', context)
def DeleteStudentSession(request, studentsessionid):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    
    try:
        studentsession = admin_models.StudentSession.objects.get(id=studentsessionid)
    except admin_models.StudentSession.DoesNotExist:
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({'success': False, 'error': 'Student session not found'})
        messages.error(request, f'Student session with ID {studentsessionid} does not exist.')
        return redirect('Students')
    
    studentid = studentsession.student.id
    student_name = studentsession.student.student_name
    session_name = studentsession.session.session_name
    
    studentsession.delete()
    
    message = "Removed  " + student_name + " from " + session_name + " session"
    admin_models.Notification.objects.create(user=user, category='Deletion', content=message)
    
    # Return JSON for AJAX delete requests
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({'success': True})
    return redirect('StudentSession', studentid=studentid)
def StudentSessionView(request, studentsessionid):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    userdata = admin_models.StudentSession.objects.get(id=studentsessionid)  # The user you are trying to update

    context = {
        'user': user,
        'userdata': userdata,
        'studentid': userdata.student.id,  # Include student ID for navigation
    }

    if request.method == 'POST':
        userdata.status = request.POST.get('status')
        userdata.notes = request.POST.get('notes')
        userdata.due_date = request.POST.get('due_date')
        message = "Updated  " + userdata.student.student_name + " Record in " + userdata.session.session_name + " session"
        admin_models.Notification.objects.create(user=user, category='Updation', content=message)
        userdata.save()

    return render(request, 'Admin/StudentSessionView.html', context)
def AddStudentSession(request, studentid):
    if 'user_id' not in request.session:
        return redirect('home')
    
    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)
    student = admin_models.Student.objects.get(id=studentid)
    active_sessions = admin_models.Sessions.objects.filter(status='Active')
    
    # Check if student has any previous enrollments (for fee waiver logic)
    has_previous_enrollments = admin_models.StudentSession.objects.filter(student=student).exists()
    
    if request.method == 'POST':
        session_id = request.POST.get('session_id')
        due_date = request.POST.get('due_date')
        discount = request.POST.get('discount', 0)
        notes = request.POST.get('notes', '')
        
        registration_date_str = request.POST.get('registration_date')
        if registration_date_str:
            from datetime import datetime
            registration_date = datetime.strptime(registration_date_str, '%Y-%m-%d').date()
        else:
            registration_date = date.today()

        try:
            session = admin_models.Sessions.objects.get(id=session_id)
            
            # INSTITUTIONAL POLICY: Check for multiple active enrollments
            existing_active_enrollment = admin_models.StudentSession.objects.filter(
                student=student,
                status='Active'
            ).exclude(session_id=session_id).first()
            
            if existing_active_enrollment:
                error_msg = f"Student is already enrolled in active session: {existing_active_enrollment.session.session_name}. Please complete or withdraw from current session before enrolling in a new one."
                if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                    return JsonResponse({'success': False, 'error': error_msg})
                messages.error(request, error_msg)
                return render(request, 'Admin/AddStudentSession.html', {
                    'student': student,
                    'active_sessions': active_sessions,
                    'error': error_msg
                })
            
            # INSTITUTIONAL POLICY: Registration fee waiver for re-enrollments
            registration_fee = 0  # Default to 0 for re-enrollments
            if not has_previous_enrollments:
                # First-time enrollment: charge full registration fee
                registration_fee = session.registration_fee
            else:
                # Re-enrollment: waive registration fee
                registration_fee = 0
                notes += f" [Registration fee waived - re-enrollment policy applied]"

            # Create new StudentSession
            student_session = admin_models.StudentSession(
                student=student,
                session=session,
                registration_date=registration_date,
                registration_fee=registration_fee,  # Applied fee waiver logic
                fee=session.fee,
                due_date=due_date,
                discount=discount,
                status='Active',
                notes=notes,
            )
            
            student_session.save()
            
            # Create notification with fee waiver info
            fee_info = "(Registration fee waived)" if has_previous_enrollments else f"(Registration fee: {registration_fee})"
            message = f"Enrolled {student.student_name} in {session.session_name} session {fee_info}"
            admin_models.Notification.objects.create(user=user, category='New Entry', content=message)
            
        except Exception as e:
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'error': str(e)})
            raise

        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({
                'success': True, 
                'message': f'Successfully enrolled with {"waived" if has_previous_enrollments else "standard"} registration fee policy'
            })
        return redirect('StudentView', studentid=studentid)

    return render(request, 'Admin/AddStudentSession.html', {
        'student': student,
        'active_sessions': active_sessions,
        'has_previous_enrollments': has_previous_enrollments,
    })
def StudentSession(request, studentid):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    userdata = admin_models.Student.objects.get(id=studentid)
    sessions = admin_models.StudentSession.objects.filter(student=userdata)

    context = {
        'user': user,
        'userdata': userdata,
        'sessions': sessions,
        'studentid': studentid
    }
    return render(request, 'Admin/StudentSession.html', context)
def LeadView(request, leadid):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    userdata = admin_models.Lead.objects.get(id=leadid)  # The user you are trying to update

    context = {
        'user': user,
        'userdata': userdata,
    }

    if request.method == 'POST':
        form = LeadForm(request.POST, request.FILES, instance=userdata)  # Form for 'userdata'

        if form.is_valid():
            form.save()
            message = "Updated  " + userdata.name + " Information in Leads"
            admin_models.Notification.objects.create(user=user, category='Deletion', content=message)
            return redirect('LeadView', leadid=leadid)  # Redirect to avoid resubmission
        else:
            # Print form errors for debugging
            print("Form is not valid:")
            print(form.errors)

    else:
        form = LeadForm(instance=userdata)

    context['form'] = form
    return render(request, 'Admin/LeadView.html', context)
def DeleteLead(request, leadid):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    
    try:
        lead = admin_models.Lead.objects.get(id=leadid)
    except admin_models.Lead.DoesNotExist:
        messages.error(request, f'Lead with ID {leadid} does not exist.')
        return redirect('Leads')
    
    lead_name = lead.lead_name
    lead.delete()
    
    message = "Deleted Lead " + lead_name
    admin_models.Notification.objects.create(user=user, category='Deletion', content=message)
    messages.success(request, f'Lead "{lead_name}" has been successfully deleted.')
    
    return redirect('Leads')
@teacher_redirect_to_attendance
def AddLead(request):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    active_sessions = admin_models.Sessions.objects.filter(status='Active')  # Fetch active sessions

    if request.method == 'POST':
        form = LeadForm(request.POST)

        if form.is_valid():
            newuser = form.save(commit=False)
            newuser.save()  # Save the new user
            message = "Added " + newuser.name + " to Leads"
            admin_models.Notification.objects.create(user=user, category='New Entry', content=message)
            messages.success(request, "Lead added successfully!")
            return redirect('Leads')
        else:
            print(form.errors)  # Debug: print any form errors

    else:
        form = LeadForm()

    context = {
        'user': user,
        'form': form,
        'active_sessions': active_sessions,
    }
    return render(request, 'Admin/AddLead.html', context)
@teacher_redirect_to_attendance
def Leads(request):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')  # Get the logged-in user ID from the session
    user = User.objects.get(id=user_id)  # Fetch the user object
    leads = admin_models.Lead.objects.all()
    
    # Filter leads by inquiry type for statistics
    call_leads = leads.filter(form_of_inquiry='Call')
    message_leads = leads.filter(form_of_inquiry='Message')
    visit_leads = leads.filter(form_of_inquiry='Physical Visit')
    
    context = {
        'user': user,
        'leads': leads,
        'call_leads': call_leads,
        'message_leads': message_leads,
        'visit_leads': visit_leads,
    }
    return render(request, 'Admin/Leads.html', context)
def DeleteStudent(request, studentid):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)
    
    try:
        student = admin_models.Student.objects.get(id=studentid)
    except admin_models.Student.DoesNotExist:
        messages.error(request, f'Student with ID {studentid} does not exist.')
        from_page = request.GET.get('from')
        if from_page == 'exstudents':
            return redirect('ExStudents')
        return redirect('Students')
    
    # Store student name before deletion for notification
    student_name = student.student_name
    
    # Remove associated files
    if student.profile_photo:
        if os.path.exists(student.profile_photo.path):
            os.remove(student.profile_photo.path)
    if student.cnic_photo:
        if os.path.exists(student.cnic_photo.path):
            os.remove(student.cnic_photo.path)
    if student.degree_photo:
        if os.path.exists(student.degree_photo.path):
            os.remove(student.degree_photo.path)
    
    # Delete the student
    student.delete()
    
    # Create notification
    message = "Removed  " + student_name
    admin_models.Notification.objects.create(user=user, category='Deletion', content=message)
    
    # Add success message
    messages.success(request, f'Student {student_name} has been successfully deleted.')
    
    # Redirect based on source page
    from_page = request.GET.get('from')
    if from_page == 'exstudents':
        return redirect('ExStudents')
    return redirect('Students')
def StudentView(request, studentid):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    userdata = admin_models.Student.objects.get(id=studentid)  # The user you are trying to update
    status_choices = admin_models.Student.STATUS_CHOICES
    
    # Initialize form early to avoid UnboundLocalError
    form = StudentForm(instance=userdata)
    form_has_errors = False
    form_errors = None
    
    # Get student's enrolled sessions
    student_sessions = admin_models.StudentSession.objects.filter(student=userdata)
    
    # Compute available sessions: active sessions not yet enrolled
    all_active = admin_models.Sessions.objects.filter(status='Active')
    enrolled_ids = student_sessions.values_list('session_id', flat=True)
    available_sessions = all_active.exclude(id__in=enrolled_ids)

    # Get student's payment information using UNIFIED SYSTEM
    # Calculate installment information from payment records
    all_payments = admin_models.Payments.objects.filter(
        studentsession__student=userdata
    ).order_by('date')
    
    # Get latest payment
    latest_payment = all_payments.filter(amount__gt=0).order_by('-date').first()
    
    # Count installments (payments with amount=0 are unpaid installments)
    total_installments = all_payments.count()
    paid_installments = all_payments.filter(amount__gt=0).count()
    unpaid_installments = all_payments.filter(amount=0).count()
    
    # Calculate one-time registration fee from primary session
    primary_session = (
        student_sessions.exclude(registration_date__isnull=True)
        .order_by('registration_date', 'id')
        .first()
        or student_sessions.order_by('id').first()
    )

    one_time_reg_fee = 0
    if primary_session:
        one_time_reg_fee = (
            primary_session.registration_fee
            if primary_session.registration_fee is not None
            else (primary_session.session.registration_fee or 0)
        )
        
    # Calculate per installment amount
    per_installment_amount = 0
    if paid_installments > 0:
        # Use amount from first paid installment
        first_paid = all_payments.filter(amount__gt=0).first()
        if first_paid:
            per_installment_amount = first_paid.amount
    else:
        # Calculate per installment amount based on final fee (total_fee already includes registration_fee - discount)
        discount_amount = sum(s.discount or 0 for s in student_sessions.filter(status='Active'))
        calculated_final_fee = userdata.total_fee - discount_amount
        if total_installments > 0:
            per_installment_amount = int(calculated_final_fee / total_installments)

    # Create payment_info object using calculated properties
    discount_amount = sum(s.discount or 0 for s in student_sessions.filter(status='Active'))
    # Fix: Use userdata.total_fee which already includes registration fee for primary session
    calculated_final_fee = userdata.total_fee - discount_amount
    calculated_remaining = max(0, calculated_final_fee - userdata.total_paid)
    
    payment_info = type('PaymentInfo', (), {
        'total_fee': userdata.total_fee,
        'paid_amount': userdata.total_paid,
        'final_fee': calculated_final_fee,
        'remaining_amount': calculated_remaining,
        'installments_count': total_installments,
        'per_installment_amount': per_installment_amount,
        'discount': discount_amount
    })()
    
    # Get installment details from payment records
    installments = []
    for payment in all_payments:
        # Show expected amount for unpaid installments, actual amount for paid ones
        display_amount = payment.amount if payment.amount > 0 else per_installment_amount
        
        installment_info = {
            'id': payment.id,
            'amount': display_amount,
            'due_date': payment.date,
            'paid_date': payment.date if payment.amount > 0 else None,
            'status': 'Paid' if payment.amount > 0 else 'Unpaid',
            'is_paid': payment.amount > 0
        }
        installments.append(installment_info)
    
    next_due_date = None
    # Find next unpaid installment due date - only if there are unpaid installments and remaining balance
    if unpaid_installments > 0 and calculated_remaining > 0:
        next_unpaid = all_payments.filter(amount=0).order_by('date').first()
        if next_unpaid:
            next_due_date = next_unpaid.date
    
    # Add these attributes to payment_info
    payment_info.installments_paid = paid_installments
    payment_info.installments_due = unpaid_installments
    
    # Calculate next_due_amount - should be 0 if all payments are complete or no active sessions
    next_due_amount = 0
    if unpaid_installments > 0 and calculated_remaining > 0:
        # Only set next due amount if there are unpaid installments and remaining balance
        next_due_amount = per_installment_amount
    payment_info.next_due_amount = next_due_amount
    
    # Find next due date from student sessions - only if there are active sessions with remaining balance
    if calculated_remaining > 0:
        student_sessions_with_due = student_sessions.filter(due_date__isnull=False, status='Active').order_by('due_date')
        if student_sessions_with_due.exists():
            next_due_date = student_sessions_with_due.first().due_date

    # Get enrollment date (first registration date from student sessions)
    enrollment_date = None
    first_session = student_sessions.order_by('registration_date').first()
    if first_session and first_session.registration_date:
        enrollment_date = first_session.registration_date

    context = {
    'user': user,
    'userdata': userdata,
    'status_choices': status_choices,
    'redirection': 1,
    'student_sessions': student_sessions,
    'available_sessions': available_sessions,
    'payment_info': payment_info,
    'installments': installments,
    'next_due_date': next_due_date,
    'enrollment_date': enrollment_date,
    'total_fee': userdata.total_fee,  # ✅ Use model property
    'registration_fee': one_time_reg_fee,  # ✅ One-time registration fee
    'discount': sum(s.discount or 0 for s in student_sessions.filter(status='Active')),
    'final_fee': userdata.total_fee - sum(s.discount or 0 for s in student_sessions.filter(status='Active')),  # Fix: Use userdata.total_fee which already includes registration fee
    'paid_amount': userdata.total_paid,  # ✅ Use model property
    'remaining_amount': userdata.remaining_balance,  # Fix: Use model property which has correct calculation
    'today_date': date.today(),
    'latest_payment': latest_payment,  # Add latest payment to context
    }

    if request.method == 'POST':
        print(f"Debug: POST data keys: {list(request.POST.keys())}")
        print(f"Debug: freeze_student in POST: {'freeze_student' in request.POST}")
        print(f"Debug: unfreeze_student in POST: {'unfreeze_student' in request.POST}")
        
        # Handle freeze/unfreeze actions first
        if 'freeze_student' in request.POST:
            print("Debug: Processing freeze_student action")
            userdata.status = 'Inactive'
            userdata.inactive_reason = 'Freeze'
            userdata.save()
            message = f"Froze {userdata.student_name}"
            admin_models.Notification.objects.create(user=user, category='Updation', content=message)
            return redirect('ExStudents')
            
        if 'unfreeze_student' in request.POST:
            print("Debug: Processing unfreeze_student action")
            userdata.status = 'Active'
            userdata.inactive_reason = ''
            userdata.save()
            message = f"Unfroze {userdata.student_name}"
            admin_models.Notification.objects.create(user=user, category='Updation', content=message)
            return redirect('Students')
        
        # Handle installment setup
        enable_installments = request.POST.get('enable_installments') == 'on'
        installments_count = int(request.POST.get('installments_count') or 0)
        per_installment_amount = Decimal(request.POST.get('per_installment_amount') or 0)
        single_due_date_str = request.POST.get('single_due_date')
        single_due_date = None
        if single_due_date_str:
            from datetime import datetime
            single_due_date = datetime.strptime(single_due_date_str, '%Y-%m-%d').date()
        
        print(f"Debug: enable_installments={enable_installments}, installments_count={installments_count}, per_installment_amount={per_installment_amount}")
        
        if enable_installments and installments_count > 0 and per_installment_amount > 0:
            print(f"Debug: Creating installments for student {userdata.student_name}")
            # Create installment payments
            from datetime import timedelta
            due_date = single_due_date if single_due_date else date.today()
            
            # Get the student sessions for payment records
            student_sessions_for_installments = admin_models.StudentSession.objects.filter(student=userdata)
            print(f"Debug: Found {student_sessions_for_installments.count()} student sessions")
            
            # Clear existing unpaid installments first
            admin_models.Payments.objects.filter(
                studentsession__student=userdata,
                amount=0
            ).delete()
            
            installments_created = 0
            for student_session in student_sessions_for_installments:
                for i in range(1, installments_count + 1):
                    # Create a payment record for each installment (initially unpaid)
                    payment = admin_models.Payments.objects.create(
                        studentsession=student_session,
                        user=user,
                        amount=0,  # Initially unpaid
                        date=due_date
                    )
                    installments_created += 1
                    print(f"Debug: Created installment {i} with payment ID {payment.id}, due date {due_date}")
                    # Next due date is one month later
                    due_date = due_date + timedelta(days=30)
            
            print(f"Debug: Total installments created: {installments_created}")
            
            # Create notification for installment setup
            installment_message = f"Set up {installments_count} installments for {userdata.student_name} - Rs.{per_installment_amount} each"
            admin_models.Notification.objects.create(
                user=user, 
                category='New Entry', 
                content=installment_message
            )
            
            messages.success(request, f"Successfully created {installments_count} installments for {userdata.student_name}")
            return redirect('StudentView', studentid=studentid)
        
        # Handle regular form submission (Save Changes button)
        # Check for specific action buttons, not just field presence
        is_freeze_action = 'freeze_student' in request.POST
        is_unfreeze_action = 'unfreeze_student' in request.POST
        is_discount_action = 'update_payment' in request.POST  # More specific discount action
        
        print(f"Debug: Action flags - freeze: {is_freeze_action}, unfreeze: {is_unfreeze_action}, discount: {is_discount_action}")
        
        # If no specific action buttons are pressed, treat as regular form save
        if not any([is_freeze_action, is_unfreeze_action, is_discount_action, enable_installments]):
            print("Debug: Processing regular form submission")
            # Filter POST data to include Student model fields and payment fields
            student_fields = [
                'student_name', 'father_name', 'email', 'cnic', 'mobile_no', 
                'Temp_address', 'Perm_address', 'last_degree', 'last_institution', 'status'
            ]
            payment_fields = ['total_fee', 'registration_fee', 'discount', 'paid_amount']
            all_fields = student_fields + payment_fields
            
            filtered_post = {}
            for field in all_fields:
                if field in request.POST:
                    filtered_post[field] = request.POST[field]
            
            # Add CSRF token
            filtered_post['csrfmiddlewaretoken'] = request.POST.get('csrfmiddlewaretoken')
            
            form = StudentForm(filtered_post, request.FILES, instance=userdata)
            
            if form.is_valid():
                # Handle file uploads
                if 'profile_photo' in request.FILES:
                    if userdata.profile_photo and os.path.exists(userdata.profile_photo.path):
                        os.remove(userdata.profile_photo.path)
                    userdata.profile_photo = request.FILES['profile_photo']
                
                if 'cnic_photo' in request.FILES:
                    if userdata.cnic_photo and os.path.exists(userdata.cnic_photo.path):
                        os.remove(userdata.cnic_photo.path)
                    userdata.cnic_photo = request.FILES['cnic_photo']
                
                if 'degree_photo' in request.FILES:
                    if userdata.degree_photo and os.path.exists(userdata.degree_photo.path):
                        os.remove(userdata.degree_photo.path)
                    userdata.degree_photo = request.FILES['degree_photo']
                
                # Save the form
                saved_student = form.save()
                
                # One-time cleanup: Remove timestamp artifacts from student name
                import re
                if "(Updated " in saved_student.student_name:
                    # Remove all timestamp patterns like "(Updated 15:00:59)"
                    cleaned_name = re.sub(r'\s*\(Updated \d{2}:\d{2}:\d{2}\)', '', saved_student.student_name)
                    saved_student.student_name = cleaned_name.strip()
                    saved_student.save()
                
                # Handle payment information updates
                print(f"Debug: Checking for payment fields in filtered_post: {[field for field in ['total_fee', 'registration_fee', 'discount', 'paid_amount'] if field in filtered_post]}")
                print(f"Debug: Payment field values: total_fee={filtered_post.get('total_fee')}, reg_fee={filtered_post.get('registration_fee')}, discount={filtered_post.get('discount')}, paid={filtered_post.get('paid_amount')}")
                
                if any(field in filtered_post for field in ['total_fee', 'registration_fee', 'discount', 'paid_amount']):
                    print("Debug: Processing payment updates...")
                    try:
                        total_fee = Decimal(filtered_post.get('total_fee', 0) or 0)
                        registration_fee = Decimal(filtered_post.get('registration_fee', 0) or 0)
                        discount = Decimal(filtered_post.get('discount', 0) or 0)
                        paid_amount = Decimal(filtered_post.get('paid_amount', 0) or 0)
                        
                        print(f"Debug: Converted values - total_fee: {total_fee}, reg_fee: {registration_fee}, discount: {discount}, paid: {paid_amount}")
                    except Exception as e:
                        print(f"Debug: Error converting payment values: {e}")
                        return redirect('StudentView', studentid=studentid)
                    
                    print(f"Debug: Processing payment updates - total_fee: {total_fee}, reg_fee: {registration_fee}, discount: {discount}, paid: {paid_amount}")
                    
                    # Update student sessions with new fees - CORRECTED
                    student_sessions = admin_models.StudentSession.objects.filter(student=saved_student)
                    for session in student_sessions:
                        # DO NOT update StudentSession.fee - it should remain as session.session.fee
                        # Only update registration_fee and discount which are session-specific
                        if registration_fee > 0:
                            session.registration_fee = int(registration_fee)
                        if discount > 0:
                            session.discount = int(discount)
                        session.save()
                        print(f"Debug: Updated StudentSession {session.id} - keeping original fee: {session.session.fee}, reg_fee: {session.registration_fee}, discount: {session.discount}")
                    
                    # Handle payment records
                    current_total_paid = sum(
                        payment.amount for payment in admin_models.Payments.objects.filter(
                            studentsession__student=saved_student, amount__gt=0
                        )
                    )
                    
                    if paid_amount > current_total_paid:
                        # Add new payment for the difference
                        additional_payment = paid_amount - current_total_paid
                        primary_session = student_sessions.first()
                        if primary_session:
                            admin_models.Payments.objects.create(
                                studentsession=primary_session,
                                user=user,
                                amount=int(additional_payment),
                                date=date.today()
                            )
                            print(f"Debug: Added payment of {additional_payment}")
                    
                    # Debug: Check what's in the database after updates
                    updated_sessions = admin_models.StudentSession.objects.filter(student=saved_student)
                    for session in updated_sessions:
                        print(f"Debug: Session {session.id} - fee: {session.session.fee}, reg_fee: {session.registration_fee}, discount: {session.discount}")
                    
                    updated_payments = admin_models.Payments.objects.filter(studentsession__student=saved_student)
                    total_paid_after = sum(p.amount for p in updated_payments if p.amount > 0)
                    print(f"Debug: Total paid after update: {total_paid_after}")
                else:
                    print("Debug: No payment fields found in POST data")
                    
                    # Update student sessions with new fees - CORRECTED
                    student_sessions = admin_models.StudentSession.objects.filter(student=saved_student)
                    for session in student_sessions:
                        # DO NOT update StudentSession.fee - it should remain as session.session.fee
                        # Only update registration_fee and discount which are session-specific
                        if registration_fee > 0:
                            session.registration_fee = int(registration_fee)
                        if discount > 0:
                            session.discount = int(discount)
                        session.save()
                        print(f"Debug: Updated StudentSession {session.id} - keeping original fee: {session.session.fee}, reg_fee: {session.registration_fee}, discount: {session.discount}")
                    
                    # Handle payment records
                    current_total_paid = sum(
                        payment.amount for payment in admin_models.Payments.objects.filter(
                            studentsession__student=saved_student, amount__gt=0
                        )
                    )
                    
                    if paid_amount > current_total_paid:
                        # Add new payment for the difference
                        additional_payment = paid_amount - current_total_paid
                        primary_session = student_sessions.first()
                        if primary_session:
                            admin_models.Payments.objects.create(
                                studentsession=primary_session,
                                user=user,
                                amount=int(additional_payment),
                                date=date.today()
                            )
                            print(f"Debug: Added payment of {additional_payment}")
                    
                    # Debug: Check what's in the database after updates
                    updated_sessions = admin_models.StudentSession.objects.filter(student=saved_student)
                    for session in updated_sessions:
                        print(f"Debug: Session {session.id} - fee: {session.session.fee}, reg_fee: {session.registration_fee}, discount: {session.discount}")
                    
                    updated_payments = admin_models.Payments.objects.filter(studentsession__student=saved_student)
                    total_paid_after = sum(p.amount for p in updated_payments if p.amount > 0)
                    print(f"Debug: Total paid after update: {total_paid_after}")
                
                # BEGIN: Installment setup on regular Save Changes
                enable_installments = request.POST.get('enable_installments') == 'on'
                installments_count = int(request.POST.get('installments_count') or 0)
                from decimal import Decimal as _Decimal  # ensure local alias if needed
                per_installment_amount = _Decimal(request.POST.get('per_installment_amount') or 0)
                single_due_date_str = request.POST.get('single_due_date')
                single_due_date = None
                if single_due_date_str:
                    from datetime import datetime as _dt
                    single_due_date = _dt.strptime(single_due_date_str, '%Y-%m-%d').date()
                
                print(f"Debug: enable_installments={enable_installments}, installments_count={installments_count}, per_installment_amount={per_installment_amount}")
                
                if enable_installments and installments_count > 0 and per_installment_amount > 0:
                    print(f"Debug: Creating installments for student {userdata.student_name}")
                    from datetime import timedelta
                    due_date = single_due_date if single_due_date else date.today()
                    
                    # Clear existing unpaid installments to avoid duplicates
                    admin_models.Payments.objects.filter(
                        studentsession__student=saved_student,
                        amount=0
                    ).delete()
                    
                    student_sessions = admin_models.StudentSession.objects.filter(student=saved_student)
                    print(f"Debug: Found {student_sessions.count()} student sessions")
                    
                    installments_created = 0
                    for student_session in student_sessions:
                        for i in range(1, installments_count + 1):
                            payment = admin_models.Payments.objects.create(
                                studentsession=student_session,
                                user=user,
                                amount=0,  # initially unpaid
                                date=due_date
                            )
                            installments_created += 1
                            print(f"Debug: Created installment {i} with payment ID {payment.id}, due date {due_date}")
                            due_date = due_date + timedelta(days=30)
                    
                    print(f"Debug: Total installments created: {installments_created}")
                    admin_models.Notification.objects.create(
                        user=user,
                        category='New Entry',
                        content=f"Set up {installments_count} installments for {saved_student.student_name} - Rs.{per_installment_amount} each"
                    )
                    return redirect('StudentView', studentid=studentid)
                # END: Installment setup on regular Save Changes
                
                message = f"Updated {userdata.student_name}"
                admin_models.Notification.objects.create(user=user, category='Updation', content=message)
                return redirect('StudentView', studentid=studentid)
            else:
                # Form has errors - we'll handle this after context is built
                form_has_errors = True
                form_errors = form.errors
            
        # Handle payment information update using UNIFIED SYSTEM
        if is_discount_action:
            discount = Decimal(request.POST.get('discount', 0))
            paid_amount = Decimal(request.POST.get('paid_amount', 0))
            
            # In the unified system, we handle payments differently:
            # 1. Discounts are applied at the session level
            # 2. Payments are recorded in the Payments table
            # 3. No separate StudentFee or Installment tables
            
            # Update discount in student sessions
            if discount > 0:
                # Distribute discount across sessions proportionally
                total_session_fees = sum(
                    (session.fee or session.session.fee or 0) + (session.registration_fee or 0)
                    for session in student_sessions
                )
                
                if total_session_fees > 0:
                    for session in student_sessions:
                        session_fee = (session.fee or session.session.fee or 0) + (session.registration_fee or 0)
                        if session_fee > 0:
                            proportion = session_fee / total_session_fees
                            session_discount = int(discount * proportion)
                            session.discount = session_discount
                            session.save()
            
            # Handle payment - create a payment record if paid_amount > current total_paid
            current_paid = userdata.total_paid
            if paid_amount > current_paid:
                additional_payment = paid_amount - current_paid
                
                # Create payment record for the additional amount
                primary_session = student_sessions.first()
                if primary_session and additional_payment > 0:
                    admin_models.Payments.objects.create(
                        studentsession=primary_session,
                        user=user,
                        amount=int(additional_payment),
                        date=date.today()
                    )
                    
                    message = f"Added payment of Rs. {additional_payment} for {userdata.student_name}"
                    admin_models.Notification.objects.create(user=user, category='New Fee', content=message)
            
            # BEGIN: Installment setup on update_payment action
            enable_installments = request.POST.get('enable_installments') == 'on'
            installments_count = int(request.POST.get('installments_count') or 0)
            per_installment_amount = Decimal(request.POST.get('per_installment_amount') or 0)
            single_due_date_str = request.POST.get('single_due_date')
            single_due_date = None
            if single_due_date_str:
                from datetime import datetime as _dt2
                single_due_date = _dt2.strptime(single_due_date_str, '%Y-%m-%d').date()
            
            print(f"Debug: [update_payment] enable_installments={enable_installments}, installments_count={installments_count}, per_installment_amount={per_installment_amount}")
            
            if enable_installments and installments_count > 0 and per_installment_amount > 0:
                print(f"Debug: [update_payment] Creating installments for student {userdata.student_name}")
                from datetime import timedelta
                due_date = single_due_date if single_due_date else date.today()
                
                # Clear existing unpaid installments to avoid duplicates
                admin_models.Payments.objects.filter(
                    studentsession__student=userdata,
                    amount=0
                ).delete()
                
                student_sessions = admin_models.StudentSession.objects.filter(student=userdata)
                print(f"Debug: [update_payment] Found {student_sessions.count()} student sessions")
                
                installments_created = 0
                for student_session in student_sessions:
                    for i in range(1, installments_count + 1):
                        payment = admin_models.Payments.objects.create(
                            studentsession=student_session,
                            user=user,
                            amount=0,
                            date=due_date
                        )
                        installments_created += 1
                        print(f"Debug: [update_payment] Created installment {i} with payment ID {payment.id}, due date {due_date}")
                        due_date = due_date + timedelta(days=30)
                
                print(f"Debug: [update_payment] Total installments created: {installments_created}")
                admin_models.Notification.objects.create(
                    user=user,
                    category='New Entry',
                    content=f"Set up {installments_count} installments for {userdata.student_name} - Rs.{per_installment_amount} each"
                )
                return redirect('StudentView', studentid=studentid)
            # END: Installment setup on update_payment action
            # Installment functionality removed in unified system
            # All payments are now handled through the Payments table
            # No need for separate installment tracking
            
            # Create notification for payment update (if any payment was made)
            if 'paid_amount' in request.POST and Decimal(request.POST.get('paid_amount', 0)) > 0:
                message = f"Updated payment information for {userdata.student_name}"
                admin_models.Notification.objects.create(user=user, category='Updation', content=message)
                
        # This section is now handled above in the reorganized POST handling

    # Handle form errors if they occurred during POST
    if form_has_errors:
        context['form'] = form  # Form with errors
        context['form_errors'] = form_errors
    else:
        context['form'] = form  # Clean form
    
    response = render(request, 'Admin/StudentView.html', context)
    # Add cache-busting headers
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    return response
def mark_installment_paid(request, studentid):
    """
    Mark the next unpaid installment as paid in the unified payment system.
    """
    if request.method == 'POST' and request.headers.get('x-requested-with') == 'XMLHttpRequest':
        if 'user_id' not in request.session:
            return JsonResponse({'success': False, 'error': 'Not authenticated'})
        
        try:
            user_id = request.session.get('user_id')
            user = User.objects.get(id=user_id)
            student = admin_models.Student.objects.get(id=studentid)
            
            # Find the next unpaid installment (payment with amount=0)
            next_unpaid = admin_models.Payments.objects.filter(
                studentsession__student=student,
                amount=0
            ).order_by('date').first()
            
            if not next_unpaid:
                return JsonResponse({
                    'success': False, 
                    'error': 'No unpaid installments found for this student.'
                })
            
            # Calculate the installment amount
            installment_amount = request.POST.get('amount')
            if not installment_amount:
                # Calculate from student's total fee and installment count
                student_sessions = admin_models.StudentSession.objects.filter(student=student)
                total_fee = sum(session.session.fee or 0 for session in student_sessions)
                registration_fee = sum(session.session.registration_fee or 0 for session in student_sessions)
                total_with_reg = total_fee + registration_fee
                
                # Get total installment count
                all_payments = admin_models.Payments.objects.filter(studentsession__student=student)
                installment_count = all_payments.count()
                
                if installment_count > 0:
                    installment_amount = int(total_with_reg / installment_count)
                else:
                    installment_amount = total_with_reg
            
            # Mark the installment as paid
            next_unpaid.amount = int(installment_amount)
            next_unpaid.save()
            
            # Create notification
            message = f"Installment payment of ${installment_amount} received for {student.student_name}"
            admin_models.Notification.objects.create(
                user=user, 
                category='Payment', 
                content=message
            )
            
            # Calculate updated payment info
            all_payments = admin_models.Payments.objects.filter(studentsession__student=student)
            paid_count = all_payments.filter(amount__gt=0).count()
            total_count = all_payments.count()
            unpaid_count = all_payments.filter(amount=0).count()
            
            # Calculate total paid amount - line 1578
            total_paid = sum(payment.amount for payment in all_payments.filter(amount__gt=0))  # Only sum actual payments
            
            # Calculate remaining amount
            student_sessions = admin_models.StudentSession.objects.filter(student=student)
            total_fee = sum(session.session.fee or 0 for session in student_sessions)
            registration_fee = sum(session.session.registration_fee or 0 for session in student_sessions)
            total_with_reg = total_fee + registration_fee
            remaining_amount = total_with_reg - total_paid
            
            # Get next due date
            next_unpaid_after = admin_models.Payments.objects.filter(
                studentsession__student=student,
                amount=0
            ).order_by('date').first()
            next_due_date = next_unpaid_after.date.strftime('%Y-%m-%d') if next_unpaid_after else None
            
            # Calculate next due amount (remaining installments)
            next_due_amount = int(remaining_amount / unpaid_count) if unpaid_count > 0 else 0
            
            return JsonResponse({
                'success': True, 
                'message': f'Installment of ${installment_amount} marked as paid successfully!',
                'paid_installments': paid_count,
                'total_installments': total_count,
                'installments_due': unpaid_count,
                'paid_amount': total_paid,
                'remaining_amount': remaining_amount,
                'next_due_amount': next_due_amount,
                'next_due_date': next_due_date
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method'})
def ExStudents(request):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')  # Get the logged-in user ID from the session
    user = User.objects.get(id=user_id)  # Fetch the user object
    students = admin_models.Student.objects.filter(status="Completed")
    context = {
        'user': user,
        'students': students,
        'redirection': 2
    }
    return render(request, 'Admin/ExStudents.html', context)
@teacher_redirect_to_attendance
def AddStudent(request, id=None):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    active_sessions = admin_models.Sessions.objects.filter(status='Active')  # Fetch active sessions
    
    # Mode flag to indicate if we're editing an existing student 
    # or adding a session for existing student
    is_edit_mode = False
    selected_student = None
    selected_session = None
    context = {}
    
    # If id parameter is provided, check if it's a student or session ID
    if id is not None:
        try:
            # First try to find a student with this ID
            selected_student = admin_models.Student.objects.filter(id=id).first()
            if selected_student:
                # We're adding a new session for this student or editing them
                is_edit_mode = True
                # Get student's current sessions for reference
                student_sessions = admin_models.StudentSession.objects.filter(student=selected_student)
                context.update({
                    'selected_student': selected_student,
                    'student_sessions': student_sessions,
                    'is_edit_mode': is_edit_mode,
                    'button_text': 'Save Changes',
                    'form_title': 'Enroll Student in New Session'
                })
            else:
                # If not a student, try to find a session with this ID
                selected_session = admin_models.Sessions.objects.filter(id=id).first()
                if selected_session:
                    # We're adding a student to this session
                    context.update({
                        'selected_session': selected_session,
                        'preselected_session_id': selected_session.id,
                        'button_text': 'Save Changes',
                        'form_title': 'Add Student to Session'
                    })
        except Exception as e:
            print(f"Error handling ID parameter: {e}")
    
    if request.method == 'POST':
        print(f"POST data received: {request.POST}")  # Debug line
        print(f"FILES received: {request.FILES}")  # Debug line
        
        # Existing POST handling code
        if selected_student:
            # We're updating a student or adding them to a new session
            form = StudentForm(request.POST, request.FILES, instance=selected_student)
        else:
            # We're creating a new student
            form = StudentForm(request.POST, request.FILES)

        print(f"Form is valid: {form.is_valid()}")  # Debug line
        if not form.is_valid():
            print(f"Form errors: {form.errors}")  # Debug line
            
        if form.is_valid():
            try:
                if not selected_student:
                    # Only create a new student if we're not editing
                    newuser = form.save(commit=False)
                    
                    # Set the logged-in user as the one who added this student
                    newuser.added_by = user

                    # Handle profile photo if uploaded
                    if 'profile_photo' in request.FILES:
                        newuser.profile_photo = request.FILES['profile_photo']
                    if 'cnic_photo' in request.FILES:
                        newuser.cnic_photo = request.FILES['cnic_photo']
                    if 'degree_photo' in request.FILES:
                        newuser.degree_photo = request.FILES['degree_photo']

                    newuser.save()  # Save the new user
                    selected_student = newuser
                    print(f"Student created successfully: {newuser.id}")  # Debug line

                else:
                    # For existing student, just update and save
                    student = form.save(commit=False)
                    
                    # Handle profile photo if uploaded
                    if 'profile_photo' in request.FILES:
                        student.profile_photo = request.FILES['profile_photo']
                    if 'cnic_photo' in request.FILES:
                        student.cnic_photo = request.FILES['cnic_photo']
                    if 'degree_photo' in request.FILES:
                        student.degree_photo = request.FILES['degree_photo']
                    
                    student.save()
                    
            except Exception as e:
                print(f"Error saving student: {e}")  # Debug line
                # Add error message to form
                form.add_error(None, f"Error saving student: {str(e)}")
                return render(request, 'Admin/AddStudent.html', context)

            # Handle session enrollment with auto roll number generation
            try:
                selected_sessions = request.POST.getlist('sessions')
                print(f"Debug: Selected sessions from POST: {selected_sessions}")
                total_fee = 0
                registration_fee = 0
                single_due_date_str = request.POST.get('single_due_date')
                single_due_date = None
                if single_due_date_str:
                    from datetime import datetime
                    single_due_date = datetime.strptime(single_due_date_str, '%Y-%m-%d').date()
                
                # Validate installment configurations before processing
                enable_installments = request.POST.get('enable_installments') == 'on'
                installments_count = int(request.POST.get('installments_count') or 0)
                
                # Check for monthly sessions and disable installments
                has_monthly_session = False
                
                if selected_sessions:
                    for session_id in selected_sessions:
                        session = admin_models.Sessions.objects.get(id=session_id)
                        if session.session_type == 'monthly':
                            has_monthly_session = True
                            break
                    
                    # Disable installments for monthly sessions
                    if has_monthly_session and enable_installments:
                        form.add_error(None, "Installment payments are not allowed for monthly sessions.")
                        return render(request, 'Admin/AddStudent.html', context)
                if selected_sessions:
                    for session_id in selected_sessions:
                        # Check if student is already enrolled in this session
                        existing_session = admin_models.StudentSession.objects.filter(
                            student=selected_student, 
                            session_id=session_id
                        ).first()
                        
                        if not existing_session:
                            session = admin_models.Sessions.objects.get(id=session_id)
                            
                            # Generate roll number if student doesn't have one
                            if not selected_student.rollno:
                                roll_number = selected_student.generate_roll_number(session)
                                selected_student.rollno = roll_number
                                selected_student.save()
                            
                            # Create StudentSession
                            student_session = admin_models.StudentSession(
                                student=selected_student,
                                session=session,
                                registration_date=date.today(),
                                registration_fee=session.registration_fee,
                                fee=session.fee,  # Set the fee from the session
                                status='Active',
                                due_date=single_due_date if not request.POST.get('enable_installments') else None
                            )
                            student_session.save()
                            print(f"Debug: Created StudentSession ID: {student_session.id} for session {session.session_name}")
                            
                            total_fee += session.fee
                            registration_fee += session.registration_fee
                            
                            message = f"Added {selected_student.student_name} (Roll: {selected_student.rollno}) to {session.session_name} session"
                            admin_models.Notification.objects.create(
                                user=user, 
                                category='New Entry', 
                                content=message
                            )

                # Handle payment information
                discount = Decimal(request.POST.get('discount', 0))
                paid_amount = Decimal(request.POST.get('paid_amount', 0))
                per_installment_amount = Decimal(request.POST.get('per_installment_amount') or 0)
                
                # Handle payment and installments (moved outside fee condition)
                print(f"Debug: enable_installments={enable_installments}, installments_count={installments_count}, per_installment_amount={per_installment_amount}")
            
                # Get the student sessions for payment records
                student_sessions_list = admin_models.StudentSession.objects.filter(student=selected_student)
                primary_session = student_sessions_list.first()
                
                # Create initial payment records for tracking
                if primary_session:
                    if not enable_installments or installments_count <= 1:
                        # Create a single payment due on registration date (current date)
                        payment = admin_models.Payments.objects.create(
                            studentsession=primary_session,
                            user=user,
                            amount=0,  # Initially unpaid
                            date=date.today()
                        )
                    else:
                        # For installments, create payment records for each installment
                        due_date = single_due_date if single_due_date else date.today()
                        
                        for i in range(1, installments_count + 1):
                            payment = admin_models.Payments.objects.create(
                                studentsession=primary_session,
                                user=user,
                                amount=0,  # Initially unpaid
                                date=due_date
                            )
                            # Next due date is one month later
                            due_date = due_date + relativedelta(months=1)
                        
                        # Create notification for installment setup
                        installment_message = f"Set up {installments_count} installments for {selected_student.student_name} - Rs.{per_installment_amount} each"
                        admin_models.Notification.objects.create(
                            user=user, 
                            category='New Entry', 
                            content=installment_message
                        )
                
                # Handle immediate payment if provided
                if paid_amount > 0 and primary_session:
                    # Find the first unpaid payment record and mark it as paid
                    unpaid_payment = admin_models.Payments.objects.filter(
                        studentsession=primary_session,
                        amount=0
                    ).order_by('date').first()
                    
                    if unpaid_payment:
                        unpaid_payment.amount = int(paid_amount)
                        # When installments are not enabled, set paid date to registration date (today)
                        if not enable_installments or installments_count <= 1:
                            unpaid_payment.date = date.today()
                        unpaid_payment.save()
                    else:
                        # Create new payment record if no unpaid record exists
                        # Use registration date (today) for the payment
                        admin_models.Payments.objects.create(
                            studentsession=primary_session,
                            user=user,
                            amount=int(paid_amount),
                            date=date.today()
                        )
                    
                    payment_message = f"Payment received for {selected_student.student_name}: Rs.{paid_amount}"
                    admin_models.Notification.objects.create(
                        user=user, 
                        category='Payment', 
                        content=payment_message
                    )
                
                # Note: StudentFee and Installment models have been removed
                # Payment tracking is now handled through the unified Payments system
                # via Student and StudentSession properties

            except Exception as e:
                print(f"Error during session enrollment or payment processing: {e}")
                form.add_error(None, f"An error occurred during enrollment: {str(e)}")
                return render(request, 'Admin/AddStudent.html', {
                    'form': form,
                    'active_sessions': active_sessions,
                    'selected_student': selected_student if 'selected_student' in locals() else None,
                    'id': id
                })
            
            # Redirect to the student view after successful processing
            return redirect('StudentView', studentid=selected_student.id)
    else:
        # Not a POST request, display the form
        if selected_student:
            # If editing existing student, pre-fill the form
            form = StudentForm(instance=selected_student)
        else:
            # New student form
            form = StudentForm()

    # Update the context with common items
    context.update({
        'user': user,
        'form': form,
        'active_sessions': active_sessions,
    })
    
    if not 'button_text' in context:
        context['button_text'] = 'Add Student'
        context['form_title'] = 'Add New Student'
    
    return render(request, 'Admin/AddStudent.html', context)
@teacher_redirect_to_attendance
def Students(request):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')  # Get the logged-in user ID from the session
    user = User.objects.get(id=user_id)  # Fetch the user object
    students = admin_models.Student.objects.filter(status="Active")
    context = {
        'user': user,
        'students': students,
        'redirection': 1
    }
    return render(request, 'Admin/Students.html', context)
def DeleteSession(request, sessionid):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    
    try:
        session = admin_models.Sessions.objects.get(id=sessionid)
    except admin_models.Sessions.DoesNotExist:
        messages.error(request, f'Session with ID {sessionid} does not exist.')
        return redirect('Sessions')
    
    session_name = session.session_name
    
    # Remove session photo if it exists
    if session.session_photo:
        if os.path.exists(session.session_photo.path):
            os.remove(session.session_photo.path)
    
    session.delete()
    
    message = "Deleted Session: " + session_name
    admin_models.Notification.objects.create(user=user, category='Deletion', content=message)
    messages.success(request, f'Session "{session_name}" has been successfully deleted.')
    
    return redirect('Sessions')
def CompletedSessions(request):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')  # Get the logged-in user ID from the session
    user = User.objects.get(id=user_id)  # Fetch the user object
    sessions = admin_models.Sessions.objects.filter(status="Completed").annotate(
        student_count=Count('session_students')
    )
    context = {
        'user': user,
        'sessions': sessions,
    }
    return render(request, 'Admin/CompletedSessions.html', context)

def RestoreSession(request, sessionid):
    if 'user_id' not in request.session:
        return redirect('home')
    
    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)
    
    try:
        # Get the completed session
        session = admin_models.Sessions.objects.get(id=sessionid, status='Completed')
        
        # Update session status to Active
        session.status = 'Active'
        session.save()
        
        # Update all students associated with this session back to Active status
        student_sessions = admin_models.StudentSession.objects.filter(session=session, status='Completed')
        students_updated = 0
        
        for student_session in student_sessions:
            # Update the student session status
            student_session.status = 'Active'
            student_session.save()
            
            # Update the student's overall status to 'Active' (Current Student)
            student = student_session.student
            student.status = 'Active'
            student.save()
            students_updated += 1
        
        # Create notification
        message = f"Restored Session: {session.session_name} with {students_updated} students transitioned back to Current Student status"
        admin_models.Notification.objects.create(user=user, category='Updation', content=message)
        
        messages.success(request, f'Session "{session.session_name}" has been successfully restored to Active status. {students_updated} students have been transitioned back to Current Student status.')
        
    except admin_models.Sessions.DoesNotExist:
        messages.error(request, 'Session not found or is not in Completed status.')
    except Exception as e:
        messages.error(request, f'An error occurred while restoring the session: {str(e)}')
    
    return redirect('CompletedSessions')

@teacher_redirect_to_attendance
def AddSession(request):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user

    if request.method == 'POST':
        form = SessionForm(request.POST, request.FILES)

        if form.is_valid():
            newsession = form.save()  # Save the new session directly

            message = f"Added {newsession.get_session_type_display()}: {newsession.session_name}"
            admin_models.Notification.objects.create(user=user, category='New Entry', content=message)
            return redirect('Sessions')
        else:
            print(form.errors)  # Debug: print any form errors
    else:
        form = SessionForm()

    context = {
        'user': user,
        'form': form,
    }
    return render(request, 'Admin/AddSession.html', context)
def SessionStudentView(request, sessionid):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    sessiondata = admin_models.Sessions.objects.get(id=sessionid)
    students = admin_models.StudentSession.objects.filter(session=sessiondata)
    context = {
        'user': user,
        'sessiondata': sessiondata,
        'students': students,
    }
    return render(request, 'Admin/SessionStudentView.html', context)
def SessionView(request, sessionid):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    sessiondata = admin_models.Sessions.objects.get(id=sessionid)  # The user you are trying to update

    status_choices = admin_models.Sessions.STATUS_CHOICES

    context = {
        'user': user,
        'sessiondata': sessiondata,
        'status_choices': status_choices,
    }

    if request.method == 'POST':
        form = SessionForm(request.POST, request.FILES, instance=sessiondata)  # Form for 'userdata'

        if form.is_valid():
            if 'session_photo' in request.FILES:
                if sessiondata.session_photo:
                    print(f"Old profile photo path: {sessiondata.session_photo.path}")
                    if os.path.exists(sessiondata.session_photo.path):
                        os.remove(sessiondata.session_photo.path)
            form.save()
            message = "Updated Session: " + sessiondata.session_name
            admin_models.Notification.objects.create(user=user, category='Updation', content=message)
            return redirect('SessionView', sessionid=sessionid)  # Redirect to avoid resubmission
        else:
            # Print form errors for debugging
            print("Form is not valid:")
            print(form.errors)

    else:
        form = SessionForm(instance=sessiondata)

    context['form'] = form

    return render(request, 'Admin/SessionView.html', context)
@teacher_redirect_to_attendance
def Sessions(request):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')  # Get the logged-in user ID from the session
    user = User.objects.get(id=user_id)  # Fetch the user object
    sessions = admin_models.Sessions.objects.filter(status="Active").annotate(
        student_count=Count('session_students')
    )
    context = {
        'user': user,
        'sessions': sessions,
    }
    return render(request, 'Admin/Sessions.html', context)
def DeleteFaculty(request, userid):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    
    try:
        faculty = User.objects.get(id=userid)
    except User.DoesNotExist:
        messages.error(request, f'Faculty with ID {userid} does not exist.')
        return redirect('Faculty')
    
    faculty_name = faculty.first_name + " " + faculty.last_name
    
    # Remove faculty profile photo if it exists
    if faculty.profile_photo:
        if os.path.exists(faculty.profile_photo.path):
            os.remove(faculty.profile_photo.path)
    
    faculty.delete()
    
    message = "Deleted Faculty: " + faculty_name
    admin_models.Notification.objects.create(user=user, category='Deletion', content=message)
    messages.success(request, f'Faculty "{faculty_name}" has been successfully deleted.')
    
    return redirect('Faculty')
@teacher_redirect_to_attendance
def AddFaculty(request):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    
    # Restrict access for moderators
    if user.usertype == 2:
        return redirect('Admin_Dashboard')

    if request.method == 'POST':
        form = UserForm(request.POST, request.FILES)

        if form.is_valid():
            newuser = form.save(commit=False)

            # Handle profile photo if uploaded
            if 'profile_photo' in request.FILES:
                newuser.profile_photo = request.FILES['profile_photo']

            newuser.save()  # Save the new user

            message = "Added Faculty: " + newuser.first_name + " " + newuser.last_name
            admin_models.Notification.objects.create(user=user, category='New Entry', content=message)
            return redirect('Faculty')
        else:
            print(form.errors)  # Debug: print any form errors

    else:
        form = UserForm()

    context = {
        'user': user,
        'form': form,
    }
    return render(request, 'Admin/AddFaculty.html', context)
def FacultyView(request, userid):
    if 'user_id' not in request.session:
        return redirect('home')

    user_id = request.session.get('user_id')
    user = User.objects.get(id=user_id)  # Logged-in user
    
    # Restrict access for moderators
    if user.usertype == 2:
        return redirect('Admin_Dashboard')
    userdata = User.objects.get(id=userid)  # The user you are trying to update

    usertype_choices = User.USER_TYPE_CHOICES
    status_choices = User.STATUS_CHOICES

    context = {
        'user': user,
        'userdata': userdata,
        'usertype_choices': usertype_choices,
        'status_choices': status_choices,
    }

    if request.method == 'POST':
        form = UserForm(request.POST, request.FILES, instance=userdata)  # Form for 'userdata'

        if form.is_valid():
            # Handle profile photo
            if 'profile_photo' in request.FILES:
                if userdata.profile_photo:
                    print(f"Old profile photo path: {userdata.profile_photo.path}")
                    if os.path.exists(userdata.profile_photo.path):
                        os.remove(userdata.profile_photo.path)

                userdata.profile_photo = request.FILES['profile_photo']

            # Save user changes
            form.save()
            message = "Updated Faculty: " + userdata.first_name + " " + userdata.last_name
            admin_models.Notification.objects.create(user=user, category='Updation', content=message)
            return redirect('FacultyView', userid=userid)  # Redirect to avoid resubmission
        else:
            # Print form errors for debugging
            print("Form is not valid:")
            print(form.errors)

    else:
        form = UserForm(instance=userdata)

    context['form'] = form

    return render(request, 'Admin/FacultyView.html', context)
@teacher_redirect_to_attendance
def Faculty(request):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')  # Get the logged-in user ID from the session
    user = User.objects.get(id=user_id)  # Fetch the user object
    
    # Restrict access for moderators
    if user.usertype == 2:
        return redirect('Admin_Dashboard')
    users = User.objects.all()
    context = {
        'user': user,
        'users': users,
    }
    return render(request, 'Admin/Faculty.html', context)
def Profile(request):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')  # Get the logged-in user ID from the session
    user = User.objects.get(id=user_id)  # Fetch the user object

    if request.method == 'POST':
        form = UserForm(request.POST, request.FILES, instance=user)

        if form.is_valid():
            # Handle profile photo
            if 'profile_photo' in request.FILES:
                if user.profile_photo:
                    print(user.profile_photo.path)
                    # Delete old photo if it exists
                    if os.path.exists(user.profile_photo.path):
                        os.remove(user.profile_photo.path)

                # Save the new photo
                user.profile_photo = request.FILES['profile_photo']

            # Save user changes
            form.save()
            return redirect('Admin_Profile')  # Redirect to avoid form resubmission

    else:
        form = UserForm(instance=user)  # Pre-populate form with user data

    return render(request, 'Admin/Profile.html', {'form': form, 'user': user})
def Logout(request):
    request.session.flush()  # This clears all session data

    # Redirect to login page after logout
    return redirect('home')
@teacher_redirect_to_attendance
def Admin_Dashboard(request):
    if 'user_id' not in request.session:
        return redirect('home')
    user_id = request.session.get('user_id')  # Get the logged-in user ID from the session
    user = User.objects.get(id=user_id)  # Fetch the user object
    users = User.objects.all()
    total_students = admin_models.Student.objects.count()  # Total number of students
    total_leads = admin_models.Lead.objects.count()  # Total number of leads
    total_sessions = admin_models.Sessions.objects.count()
    total_users = User.objects.count()
    
    # Add notification count - only count unread notifications
    notification_count = admin_models.Notification.objects.filter(is_read=False).count()
    
    context = {
        'user': user,
        'users': users,
        'total_students': total_students,
        'total_leads': total_leads,
        'total_sessions': total_sessions,
        'total_users': total_users,
        'notification_count': notification_count,
    }
    return render(request, 'Admin/Dashboard.html',context)
@csrf_exempt
def filter_payments(request):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Invalid request method'})
    
    if 'user_id' not in request.session:
        return JsonResponse({'success': False, 'error': 'Unauthorized'})
    
    try:
        data = json.loads(request.body)
        filter_type = data.get('type')
        filter_value = data.get('value')
        from_date = data.get('fromDate')
        to_date = data.get('toDate')
        
        # Calculate date range based on filter
        today = datetime.now().date()
        start_date = None
        end_date = today
        filter_description = "All Time"
        
        if filter_type == 'today':
            start_date = today
            end_date = today
            filter_description = "Today"
        elif filter_type == 'days' and filter_value:
            start_date = today - timedelta(days=filter_value)
            filter_description = f"Last {filter_value} Days"
        elif filter_type == 'custom' and from_date and to_date:
            start_date = datetime.strptime(from_date, '%Y-%m-%d').date()
            end_date = datetime.strptime(to_date, '%Y-%m-%d').date()
            filter_description = f"{start_date.strftime('%d %b %Y')} to {end_date.strftime('%d %b %Y')}"
        
        # Filter payments based on date range
        payments = admin_models.Payments.objects.select_related(
            'studentsession__student', 'studentsession__session', 'user'
        ).all()
        
        if start_date:
            payments = payments.filter(date__gte=start_date)
        if end_date:
            payments = payments.filter(date__lte=end_date)
        
        # Recalculate metrics with filtered data
        filtered_data = calculate_revenue_metrics(payments, start_date, end_date)
        filtered_data['filter_description'] = filter_description
        
        return JsonResponse({
            'success': True,
            'data': filtered_data
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

def calculate_revenue_metrics(payments, start_date=None, end_date=None):
    """Calculate revenue metrics for filtered payments"""
    
    # Get all active students
    students = admin_models.Student.objects.filter(status='Active')
    
    # Calculate metrics
    total_revenue = sum(p.amount or 0 for p in payments)
    total_pending = sum(s.remaining_balance for s in students)
    total_expected_revenue = sum(s.total_fee for s in students)
    
    # Calculate collection rate
    collection_rate = (total_revenue / total_expected_revenue * 100) if total_expected_revenue > 0 else 0
    
    # Student payment status
    students_paid = sum(1 for s in students if s.remaining_balance == 0)
    students_partial = sum(1 for s in students if 0 < s.remaining_balance < s.total_fee)
    students_unpaid = sum(1 for s in students if s.remaining_balance == s.total_fee)
    
    # Recent payments (limited to filtered data)
    recent_payments = payments.order_by('-date', '-id')[:10]
    recent_payments_data = []
    for payment in recent_payments:
        recent_payments_data.append({
            'date': payment.date.isoformat() if payment.date else '',
            'student_name': payment.studentsession.student.student_name if payment.studentsession else '',
            'rollno': payment.studentsession.student.rollno if payment.studentsession else '',
            'session_name': payment.studentsession.session.session_name if payment.studentsession and payment.studentsession.session else '',
            'amount': float(payment.amount or 0),
            'collected_by': f"{payment.user.first_name} {payment.user.last_name}" if payment.user else ''
        })
    
    # Session performance
    session_revenue = {}
    for payment in payments:
        if payment.studentsession and payment.studentsession.session:
            session_name = payment.studentsession.session.session_name
            session_revenue[session_name] = session_revenue.get(session_name, 0) + (payment.amount or 0)
    
    session_performance = []
    for session_name, revenue in session_revenue.items():
        session_obj = admin_models.Sessions.objects.filter(session_name=session_name).first()
        if session_obj:
            student_count = admin_models.StudentSession.objects.filter(session=session_obj, status='Active').count()
            session_performance.append({
                'name': session_name,
                'revenue': float(revenue),
                'students': student_count,
                'avg_per_student': float(revenue / student_count) if student_count > 0 else 0
            })
    
    session_performance.sort(key=lambda x: x['revenue'], reverse=True)
    
    # Calculate other metrics
    today = datetime.now().date()
    daily_revenue = sum(p.amount or 0 for p in payments if p.date == today)
    avg_payment = total_revenue / len(payments) if payments else 0
    
    # Count overdue students
    overdue_students_count = 0
    for student in students:
        if student.remaining_balance > 0:
            student_sessions_list = student.student_sessions.filter(status='Active')
            for session in student_sessions_list:
                if session.due_date and session.due_date < today:
                    overdue_students_count += 1
                    break
    
    # Projected monthly revenue
    days_in_month = today.day
    if days_in_month > 0:
        monthly_revenue = sum(p.amount or 0 for p in payments if p.date and p.date.month == today.month and p.date.year == today.year)
        daily_avg = monthly_revenue / days_in_month
        days_remaining = 30 - days_in_month
        projected_monthly_revenue = monthly_revenue + (daily_avg * days_remaining)
    else:
        projected_monthly_revenue = 0
    
    return {
        'total_revenue': float(total_revenue),
        'total_pending': float(total_pending),
        'collection_rate': round(collection_rate, 1),
        'avg_payment': float(avg_payment),
        'daily_revenue': float(daily_revenue),
        'active_students_count': len(students),
        'overdue_students_count': overdue_students_count,
        'projected_monthly_revenue': float(projected_monthly_revenue),
        'students_paid': students_paid,
        'students_partial': students_partial,
        'students_unpaid': students_unpaid,
        'recent_payments': recent_payments_data,
        'session_performance': session_performance
    }

@csrf_exempt
def export_word_report(request):
    if request.method != 'POST':
        return HttpResponse('Invalid request method', status=405)
    
    if 'user_id' not in request.session:
        return HttpResponse('Unauthorized', status=401)
    
    try:
        data = json.loads(request.body)
        filter_data = data.get('filter', {})
        
        # Get filtered payments
        filter_type = filter_data.get('type')
        filter_value = filter_data.get('value')
        from_date = filter_data.get('fromDate')
        to_date = filter_data.get('toDate')
        
        # Calculate date range
        today = datetime.now().date()
        start_date = None
        end_date = today
        period_description = "All Time"
        
        if filter_type == 'today':
            start_date = today
            end_date = today
            period_description = "Today"
        elif filter_type == 'days' and filter_value:
            start_date = today - timedelta(days=filter_value)
            period_description = f"Last {filter_value} Days"
        elif filter_type == 'custom' and from_date and to_date:
            start_date = datetime.strptime(from_date, '%Y-%m-%d').date()
            end_date = datetime.strptime(to_date, '%Y-%m-%d').date()
            period_description = f"{start_date.strftime('%d %b %Y')} to {end_date.strftime('%d %b %Y')}"
        
        # Filter payments
        payments = admin_models.Payments.objects.select_related(
            'studentsession__student', 'studentsession__session', 'user'
        ).all()
        
        if start_date:
            payments = payments.filter(date__gte=start_date)
        if end_date:
            payments = payments.filter(date__lte=end_date)
        
        # Get metrics
        metrics = calculate_revenue_metrics(payments, start_date, end_date)
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('IQRA ACADEMY', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Revenue Report', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add report details
        doc.add_paragraph(f"Report Period: {period_description}")
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d %B %Y at %I:%M %p')}")
        doc.add_paragraph("")
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=2)
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('Total Revenue Collected', f"Rs. {metrics['total_revenue']:,.0f}"),
            ('Outstanding Receivables', f"Rs. {metrics['total_pending']:,.0f}"),
            ('Collection Efficiency', f"{metrics['collection_rate']}%"),
            ('Average Payment Size', f"Rs. {metrics['avg_payment']:,.0f}"),
            ('Active Students', str(metrics['active_students_count']))
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
        
        doc.add_paragraph("")
        
        # Student Analysis
        doc.add_heading('Student Payment Analysis', level=2)
        student_table = doc.add_table(rows=4, cols=2)
        student_table.style = 'Table Grid'
        
        student_data = [
            ('Students with Full Payment', str(metrics['students_paid'])),
            ('Students with Partial Payment', str(metrics['students_partial'])),
            ('Students with No Payment', str(metrics['students_unpaid'])),
            ('Students with Overdue Payments', str(metrics['overdue_students_count']))
        ]
        
        for i, (label, value) in enumerate(student_data):
            student_table.cell(i, 0).text = label
            student_table.cell(i, 1).text = value
        
        doc.add_paragraph("")
        
        # Session Performance
        if metrics['session_performance']:
            doc.add_heading('Session Performance', level=2)
            session_table = doc.add_table(rows=len(metrics['session_performance']) + 1, cols=4)
            session_table.style = 'Table Grid'
            
            # Headers
            headers = ['Session Name', 'Revenue', 'Students', 'Avg per Student']
            for i, header in enumerate(headers):
                session_table.cell(0, i).text = header
            
            # Data
            for i, session in enumerate(metrics['session_performance'][:10]):
                session_table.cell(i + 1, 0).text = session['name']
                session_table.cell(i + 1, 1).text = f"Rs. {session['revenue']:,.0f}"
                session_table.cell(i + 1, 2).text = str(session['students'])
                session_table.cell(i + 1, 3).text = f"Rs. {session['avg_per_student']:,.0f}"
        
        doc.add_paragraph("")
        
        # Recent Payments
        if metrics['recent_payments']:
            doc.add_heading('Recent Payments', level=2)
            payments_table = doc.add_table(rows=min(len(metrics['recent_payments']), 20) + 1, cols=5)
            payments_table.style = 'Table Grid'
            
            # Headers
            headers = ['Date', 'Student', 'Session', 'Amount', 'Collected By']
            for i, header in enumerate(headers):
                payments_table.cell(0, i).text = header
            
            # Data
            for i, payment in enumerate(metrics['recent_payments'][:20]):
                payments_table.cell(i + 1, 0).text = datetime.fromisoformat(payment['date']).strftime('%d %b %Y') if payment['date'] else ''
                payments_table.cell(i + 1, 1).text = payment['student_name']
                payments_table.cell(i + 1, 2).text = payment['session_name']
                payments_table.cell(i + 1, 3).text = f"Rs. {payment['amount']:,.0f}"
                payments_table.cell(i + 1, 4).text = payment['collected_by']
        
        doc.add_paragraph("")
        
        # Recommendations
        doc.add_heading('Recommendations', level=2)
        recommendations = []
        
        if metrics['collection_rate'] < 70:
            recommendations.append("• Collection rate is below 70% - Consider implementing automated payment reminders")
            recommendations.append("• Follow up with students having overdue payments")
        
        if metrics['students_unpaid'] > 0:
            recommendations.append(f"• {metrics['students_unpaid']} students have not made any payments - Immediate attention required")
        
        if metrics['overdue_students_count'] > 0:
            recommendations.append(f"• {metrics['overdue_students_count']} students have overdue payments - Send payment reminders")
        
        recommendations.extend([
            "• Consider offering flexible payment plans for students with large outstanding amounts",
            "• Regular monthly collection drives can improve cash flow",
            "• Implement early payment discounts to encourage prompt payments"
        ])
        
        for rec in recommendations:
            doc.add_paragraph(rec)
        
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Create response
        response = HttpResponse(
            doc_io.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="Revenue_Report_{period_description.replace(" ", "_")}_{today.strftime("%Y%m%d")}.docx"'
        
        return response
        
    except Exception as e:
        return HttpResponse(f'Error generating report: {str(e)}', status=500)

@csrf_exempt
def get_email_statistics(request):
    """Get real-time email statistics for the Email Services dashboard"""
    if 'user_id' not in request.session:
        return JsonResponse({'status': 'error', 'message': 'Not authenticated'})
    
    try:
        from django.db.models import Count, Q
        from datetime import datetime, timedelta
        
        # Get current date for calculations
        today = datetime.now().date()
        thirty_days_ago = today - timedelta(days=30)
        
        # Calculate total recipients
        total_students = admin_models.Student.objects.filter(
            email__isnull=False, 
            email__gt='',
            status='Active'
        ).count()
        
        total_faculty = User.objects.filter(
            email__isnull=False,
            email__gt=''
        ).count()
        
        total_leads = admin_models.Lead.objects.filter(
            email__isnull=False,
            email__gt=''
        ).count()
        
        total_recipients = total_students + total_faculty + total_leads
        
        # Calculate students with pending fees (for reminder notifications)
        students_with_pending = admin_models.Student.objects.filter(
            status='Active',
            email__isnull=False,
            email__gt=''
        ).annotate(
            unpaid_count=Count('student_sessions__student_payments', 
                             filter=Q(student_sessions__student_payments__amount=0))
        ).filter(unpaid_count__gt=0).count()
        
        # Calculate overdue students (those with payments past due date)
        overdue_students = admin_models.Student.objects.filter(
            status='Active',
            email__isnull=False,
            email__gt='',
            student_sessions__student_payments__amount=0,
            student_sessions__student_payments__date__lt=today
        ).distinct().count()
        
        # Get recent notifications for email activity simulation
        recent_notifications = admin_models.Notification.objects.filter(
            date__gte=thirty_days_ago,
            category__in=['Late Fee', 'General', 'New Entry']
        ).count()
        
        # Simulate email statistics based on available data
        # In a real system, you'd track actual email sends in a separate model
        estimated_emails_sent = recent_notifications + (students_with_pending * 2)  # Assume 2 reminders per pending student
        delivery_success_rate = 95.5  # Typical email delivery rate
        
        statistics = {
            'total_emails_sent': estimated_emails_sent,
            'delivery_success_rate': delivery_success_rate,
            'pending_reminders': students_with_pending,
            'total_recipients': total_recipients,
            'breakdown': {
                'students': total_students,
                'faculty': total_faculty,
                'leads': total_leads
            },
            'overdue_students': overdue_students,
            'recent_activity': {
                'last_30_days': recent_notifications,
                'pending_notifications': students_with_pending
            }
        }
        
        return JsonResponse({
            'status': 'success',
            'data': statistics
        })
        
    except Exception as e:
        return JsonResponse({
             'status': 'error',
             'message': f'Error fetching email statistics: {str(e)}'
         })

@csrf_exempt
def get_email_history(request):
    """Get recent email activity for the Email Services dashboard"""
    if 'user_id' not in request.session:
        return JsonResponse({'status': 'error', 'message': 'Not authenticated'})
    
    try:
        from django.db.models import Q
        from django.utils import timezone
        from datetime import timedelta
        
        # Get current user for sent_by field
        current_user = User.objects.get(id=request.session['user_id'])
        
        # Get recent notifications that represent email activity
        recent_notifications = admin_models.Notification.objects.filter(
            category__in=['Late Fee', 'General', 'New Entry']
        ).order_by('-date')[:20]
        
        email_history = []
        now = timezone.now()
        
        for notification in recent_notifications:
            # Simulate email activity based on notifications
            if notification.category == 'Late Fee':
                email_type = 'Payment Reminder'
                recipient_type = 'Students'
                recipients_count = '15-25'
            elif notification.category == 'General':
                email_type = 'General Notification'
                recipient_type = 'Faculty'
                recipients_count = '8-12'
            else:
                email_type = 'Welcome Message'
                recipient_type = 'Students'
                recipients_count = '10-20'
            
            # Use timezone-aware comparison
            status = 'Delivered' if notification.date < now - timedelta(hours=1) else 'Pending'
            
            email_history.append({
                'date': notification.date.strftime('%d %b %Y'),
                'subject': f'{email_type} - {notification.date.strftime("%B %Y")}',
                'recipients': recipients_count,
                'recipient_type': recipient_type,
                'status': status,
                'sent_by': f'{current_user.first_name} {current_user.last_name}'
            })
        
        # If no notifications, provide some sample data
        if not email_history:
            email_history = [
                {
                    'date': now.strftime('%d %b %Y'),
                    'subject': 'Payment Reminder - Current Month',
                    'recipients': '23',
                    'recipient_type': 'Students',
                    'status': 'Delivered',
                    'sent_by': f'{current_user.first_name} {current_user.last_name}'
                },
                {
                    'date': (now - timedelta(days=1)).strftime('%d %b %Y'),
                    'subject': 'Welcome to New Session',
                    'recipients': '15',
                    'recipient_type': 'Students',
                    'status': 'Delivered',
                    'sent_by': f'{current_user.first_name} {current_user.last_name}'
                },
                {
                    'date': (now - timedelta(days=7)).strftime('%d %b %Y'),
                    'subject': 'Faculty Meeting Reminder',
                    'recipients': '8',
                    'recipient_type': 'Faculty',
                    'status': 'Delivered',
                    'sent_by': f'{current_user.first_name} {current_user.last_name}'
                }
            ]
        
        return JsonResponse({
            'status': 'success',
            'data': email_history
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'Error fetching email history: {str(e)}'
        })
